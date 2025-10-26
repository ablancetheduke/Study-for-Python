# -*- coding: utf-8 -*-
from flask import Flask, render_template, jsonify, send_from_directory, request, redirect, url_for, current_app
from flask_cors import CORS
from flask_socketio import SocketIO, emit, join_room, leave_room
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from pymongo import MongoClient
from werkzeug.utils import secure_filename
from datetime import datetime, UTC, timedelta
from pathlib import Path
import os
from dotenv import load_dotenv
from flask import session, request, redirect, url_for, render_template
import requests

# 加载环境变量
load_dotenv()
from bson import ObjectId
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from io import BytesIO
import base64
from PyPDF2 import PdfReader
from docx import Document
import re
import json
from collections import Counter
import bcrypt
import uuid
import hmac
import hashlib
import websocket
import threading
import time
import ssl
# =========================
# Flask 初始化
# =========================
app = Flask(
    __name__,
    static_folder='app/static',
    template_folder='app/templates'
)
CORS(app)

# 从环境变量获取配置
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'default-secret-key-change-in-production').encode()

# JWT配置
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', 'default-jwt-secret-key-change-in-production')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=int(os.getenv('JWT_ACCESS_TOKEN_EXPIRES_HOURS', 24)))

# 初始化扩展
jwt = JWTManager(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# =========================
# MongoDB
# =========================
mongodb_uri = os.getenv('MONGODB_URI', 'mongodb://localhost:27017/')
client = MongoClient(mongodb_uri)

# 主库（国家主数据统一存放）
MASTER_DB_NAME = "countriesDB"
MASTER_DB = client[MASTER_DB_NAME]
MASTER_COUNTRIES = MASTER_DB["countries_lc"]

# 用户管理集合
USERS_COLLECTION = MASTER_DB["users"]
USER_SESSIONS_COLLECTION = MASTER_DB["user_sessions"]
MEETING_ROOMS_COLLECTION = MASTER_DB["meeting_rooms"]

import re

# =========================
# 用户模型
# =========================
class User(UserMixin):
    def __init__(self, user_id, username, email, role=None):
        self.id = user_id
        self.username = username
        self.email = email
        self.role = role

@login_manager.user_loader
def load_user(user_id):
    """加载用户"""
    user_data = USERS_COLLECTION.find_one({"_id": ObjectId(user_id)})
    if user_data:
        return User(
            user_id=str(user_data["_id"]),
            username=user_data["username"],
            email=user_data["email"],
            role=user_data.get("role")
        )
    return None

def get_cols_by_session(session_id: str):
    """根据五位数会议编号选择独立数据库，否则回落到主库。
    返回各功能集合的句柄。
    """
    if session_id and re.fullmatch(r"\d{5}", str(session_id)):
        db = client[f"{MASTER_DB_NAME}_{session_id}"]
    else:
        db = MASTER_DB
    return {
        "db": db,
        "countries": MASTER_COUNTRIES,  # 国家主数据仍走主库
        "settings": db["meeting_settings"],
        "rollcall": db["rollcall"],
        "submissions": db["submissions"],
        "temp_files": db["temp_files"],
        "file_assignments": db["file_assignments"],
        "vote_files": db["vote_files"],
        "vote_results": db["vote_results"],
        "file_vote_details": db["file_vote_details"],
        "declarations": db["declarations"],
        "chairman_settings": db["chairman_settings"],
        "voting_mechanisms": db["voting_mechanisms"],
    }

# 为兼容现有代码：保留默认集合指向主库
col_countries = MASTER_COUNTRIES
col_settings  = MASTER_DB["meeting_settings"]
col_rollcall = MASTER_DB["rollcall"]
col_submissions = MASTER_DB["submissions"]
col_temp_files = MASTER_DB["temp_files"]
# 工具方法：确保会话在全局集合存在
def ensure_session_in_global(session_id: str) -> bool:
    try:
        if not session_id:
            return False
        # 已存在则返回
        if col_settings.find_one({"session_id": session_id}):
            return True
        # 尝试从会期专属库同步
        cols = get_cols_by_session(session_id)
        sdoc = cols["settings"].find_one({"session_id": session_id})
        if not sdoc:
            return False
        payload = {
            "session_id": sdoc.get("session_id"),
            "committee_name": sdoc.get("committee_name"),
            "agenda": sdoc.get("agenda"),
            "created_by": sdoc.get("created_by"),
            "created_at": sdoc.get("created_at"),
            "status": sdoc.get("status", "active"),
            "participants": sdoc.get("participants", []),
            "meeting_state": sdoc.get("meeting_state", {}),
            "chairman_controls": sdoc.get("chairman_controls", {}),
        }
        col_settings.insert_one(payload)
        return True
    except Exception as _e:
        print(f"ensure_session_in_global 失败: {_e}")
        return False

col_file_assignments = MASTER_DB["file_assignments"]
col_vote_files = MASTER_DB["vote_files"]

# 索引：一国在同一会期只能提交一次（主库默认）
col_submissions.create_index([("country_id", 1), ("session_id", 1)], unique=True, name="uniq_country_session")

def ensure_indexes(target_db):
    try:
        target_db["submissions"].create_index([("country_id", 1), ("session_id", 1)], unique=True, name="uniq_country_session")
    except Exception:
        pass

# =========================
# 用户认证API
# =========================

@app.route('/api/auth/register', methods=['POST'])
def api_register():
    """用户注册"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        email = data.get('email', '').strip()
        password = data.get('password', '').strip()
        
        if not username or not email or not password:
            return jsonify({
                "code": 400,
                "message": "用户名、邮箱和密码不能为空"
            }), 400
        
        # 检查用户名是否已存在
        if USERS_COLLECTION.find_one({"username": username}):
            return jsonify({
                "code": 409,
                "message": "用户名已存在"
            }), 409
        
        # 检查邮箱是否已存在
        if USERS_COLLECTION.find_one({"email": email}):
            return jsonify({
                "code": 409,
                "message": "邮箱已被注册"
            }), 409
        
        # 加密密码
        password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        
        # 创建用户
        user_data = {
            "username": username,
            "email": email,
            "password_hash": password_hash,
            "role": "participant",  # 默认为参与国
            "created_at": datetime.now(UTC).isoformat() + "Z",
            "last_login": None,
            "status": "active"
        }
        
        result = USERS_COLLECTION.insert_one(user_data)
        user_id = str(result.inserted_id)
        
        # 创建JWT token
        access_token = create_access_token(identity=user_id)
        
        return jsonify({
            "code": 200,
            "message": "注册成功",
            "data": {
                "user_id": user_id,
                "username": username,
                "email": email,
                "access_token": access_token
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"注册失败: {str(e)}"
        }), 500

@app.route('/api/auth/login', methods=['POST'])
def api_login():
    """用户登录"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        password = data.get('password', '').strip()
        
        if not username or not password:
            return jsonify({
                "code": 400,
                "message": "用户名和密码不能为空"
            }), 400
        
        # 查找用户
        user_data = USERS_COLLECTION.find_one({"username": username})
        if not user_data:
            return jsonify({
                "code": 401,
                "message": "用户名或密码错误"
            }), 401
        
        # 验证密码
        if not bcrypt.checkpw(password.encode('utf-8'), user_data['password_hash'].encode('utf-8')):
            return jsonify({
                "code": 401,
                "message": "用户名或密码错误"
            }), 401
        
        # 更新最后登录时间
        USERS_COLLECTION.update_one(
            {"_id": user_data["_id"]},
            {"$set": {"last_login": datetime.now(UTC).isoformat() + "Z"}}
        )
        
        # 创建JWT token
        access_token = create_access_token(identity=str(user_data["_id"]))
        
        return jsonify({
            "code": 200,
            "message": "登录成功",
            "data": {
                "user_id": str(user_data["_id"]),
                "username": user_data["username"],
                "email": user_data["email"],
                "role": user_data.get("role", "participant"),
                "access_token": access_token
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"登录失败: {str(e)}"
        }), 500

@app.route('/api/auth/logout', methods=['POST'])
@jwt_required()
def api_logout():
    """用户登出"""
    try:
        # 这里可以添加token黑名单逻辑
        return jsonify({
            "code": 200,
            "message": "登出成功"
        })
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"登出失败: {str(e)}"
        }), 500

@app.route('/api/auth/profile', methods=['GET'])
@jwt_required()
def api_get_profile():
    """获取用户信息"""
    try:
        user_id = get_jwt_identity()
        user_data = USERS_COLLECTION.find_one({"_id": ObjectId(user_id)})
        
        if not user_data:
            return jsonify({
                "code": 404,
                "message": "用户不存在"
            }), 404
        
        return jsonify({
            "code": 200,
            "data": {
                "user_id": str(user_data["_id"]),
                "username": user_data["username"],
                "email": user_data["email"],
                "role": user_data.get("role", "participant"),
                "created_at": user_data.get("created_at"),
                "last_login": user_data.get("last_login")
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"获取用户信息失败: {str(e)}"
        }), 500

# =========================
# 国家选择管理API
# =========================

@app.route('/api/save_country_selection', methods=['POST', 'GET'])
def api_save_country_selection():
    """保存国家选择到数据库，或在 GET 时检查是否已保存"""
    try:
        # GET: 检查是否已存在相同的国家选择（用于前端回填判断）
        if request.method == 'GET':
            session_id = request.args.get('session_id')
            country_id = request.args.get('country_id')

            if not session_id or not country_id:
                return jsonify({
                    'code': 400,
                    'message': '缺少必要参数'
                }), 400

            session_info = col_settings.find_one({"session_id": session_id})
            if not session_info:
                return jsonify({
                    'code': 404,
                    'message': '会议不存在'
                }), 404

            exists = col_settings.find_one({
                "session_id": session_id,
                "participants.country_id": country_id
            })

            if exists:
                return jsonify({
                    'code': 200,
                    'message': '该国家已经加入会议'
                })
            else:
                return jsonify({
                    'code': 404,
                    'message': '该国家尚未加入会议'
                }), 404

        # POST: 保存国家选择
        data = request.get_json()
        print(f"收到保存国家选择请求: {data}")
        
        session_id = data.get('session_id')
        country_id = data.get('country_id')
        country_name = data.get('country_name')
        country_flag = data.get('country_flag')
        
        print(f"解析参数: session_id={session_id}, country_id={country_id}, country_name={country_name}")
        
        if not all([session_id, country_id, country_name]):
            print("缺少必要参数")
            return jsonify({
                'code': 400,
                'message': '缺少必要参数'
            }), 400
        
        # 检查会议是否存在（全局集合）
        session_info = col_settings.find_one({"session_id": session_id})
        print(f"查找会议(全局): {session_info}")
        
        if not session_info:
            # 兜底：尝试从会期专属库同步一份到全局集合
            synced = ensure_session_in_global(session_id)
            if synced:
                session_info = col_settings.find_one({"session_id": session_id})
                print(f"同步后查找会议(全局): {session_info}")
            else:
                print("会议不存在（全局与专属库均未找到）")
                return jsonify({
                    'code': 404,
                    'message': '会议不存在'
                }), 404
        
        # 检查是否已经存在相同的国家选择
        existing_participant = col_settings.find_one({
            "session_id": session_id,
            "participants.country_id": country_id
        })
        print(f"检查现有参与国: {existing_participant}")
        
        if existing_participant:
            print("该国家已经加入会议")
            return jsonify({
                'code': 409,
                'message': '该国家已经加入会议'
            }), 409
        
        # 保存国家选择到会议设置
        participant_data = {
            "country_id": country_id,
            "country_name": country_name,
            "country_flag": country_flag,
            "joined_at": datetime.now(UTC).isoformat() + "Z",
            "status": "active"
        }
        print(f"准备保存参与国数据: {participant_data}")
        
        result = col_settings.update_one(
            {"session_id": session_id},
            {
                "$push": {
                    "participants": participant_data
                }
            }
        )
        
        print(f"更新结果: matched_count={result.matched_count}, modified_count={result.modified_count}")
        
        if result.matched_count == 0:
            print("会议不存在，无法更新")
            return jsonify({
                'code': 404,
                'message': '会议不存在'
            }), 404
        
        return jsonify({
            'code': 200,
            'message': '国家选择已保存',
            'data': {
                'session_id': session_id,
                'country_id': country_id,
                'country_name': country_name
            }
        })
        
    except Exception as e:
        print(f"保存国家选择失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'code': 500,
            'message': f'服务器内部错误: {str(e)}'
        }), 500

@app.route('/api/meeting/participants', methods=['GET'])
def api_get_meeting_participants():
    """获取会议的参与国列表"""
    try:
        session_id = request.args.get('session_id')
        
        if not session_id:
            return jsonify({
                'code': 400,
                'message': '缺少会议编号参数'
            }), 400
        
        # 获取会议信息
        session_info = col_settings.find_one({"session_id": session_id})
        if not session_info:
            return jsonify({
                'code': 404,
                'message': '会议不存在'
            }), 404
        
        # 获取参与国列表
        participants = session_info.get('participants', [])
        
        return jsonify({
            'code': 200,
            'message': '获取参与国列表成功',
            'data': participants
        })
        
    except Exception as e:
        print(f"获取参与国列表失败: {e}")
        return jsonify({
            'code': 500,
            'message': '服务器内部错误'
        }), 500

# =========================
# 会议房间管理API
# =========================

@app.route('/api/rooms/create', methods=['POST'])
@jwt_required()
def api_create_room():
    """创建会议房间"""
    try:
        data = request.get_json()
        session_id = data.get('session_id', '').strip()
        committee_name = data.get('committee_name', '').strip()
        agenda = data.get('agenda', '').strip()
        max_participants = data.get('max_participants', 50)
        
        if not session_id or not committee_name:
            return jsonify({
                "code": 400,
                "message": "会议编号和委员会名称不能为空"
            }), 400
        
        user_id = get_jwt_identity()
        
        # 检查房间是否已存在
        existing_room = MEETING_ROOMS_COLLECTION.find_one({"session_id": session_id})
        if existing_room:
            return jsonify({
                "code": 409,
                "message": "会议编号已存在"
            }), 409
        
        # 创建会议房间
        room_data = {
            "room_id": str(uuid.uuid4()),
            "session_id": session_id,
            "committee_name": committee_name,
            "agenda": agenda,
            "max_participants": max_participants,
            "current_participants": 0,
            "room_status": "waiting",  # waiting, active, completed
            "created_by": user_id,
            "created_at": datetime.now(UTC).isoformat() + "Z",
            "participants": []
        }
        
        result = MEETING_ROOMS_COLLECTION.insert_one(room_data)
        
        return jsonify({
            "code": 200,
            "message": "会议房间创建成功",
            "data": {
                "room_id": room_data["room_id"],
                "session_id": session_id,
                "committee_name": committee_name,
                "agenda": agenda
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"创建会议房间失败: {str(e)}"
        }), 500

@app.route('/api/rooms/list', methods=['GET'])
def api_list_rooms():
    """获取可用会议房间列表"""
    try:
        rooms = list(MEETING_ROOMS_COLLECTION.find(
            {"room_status": {"$in": ["waiting", "active"]}},
            {"password_hash": 0}  # 不返回密码
        ).sort("created_at", -1))
        
        # 格式化数据
        room_list = []
        for room in rooms:
            room_list.append({
                "room_id": room["room_id"],
                "session_id": room["session_id"],
                "committee_name": room["committee_name"],
                "agenda": room["agenda"],
                "max_participants": room["max_participants"],
                "current_participants": room["current_participants"],
                "room_status": room["room_status"],
                "created_at": room["created_at"]
            })
        
        return jsonify({
            "code": 200,
            "data": room_list
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"获取房间列表失败: {str(e)}"
        }), 500

@app.route('/api/rooms/join', methods=['POST'])
@jwt_required()
def api_join_room():
    """加入会议房间"""
    try:
        data = request.get_json()
        room_id = data.get('room_id', '').strip()
        role = data.get('role', 'participant')  # chairman, participant
        country_id = data.get('country_id', '')  # 如果是参与国
        
        if not room_id:
            return jsonify({
                "code": 400,
                "message": "房间ID不能为空"
            }), 400
        
        user_id = get_jwt_identity()
        
        # 查找房间
        room = MEETING_ROOMS_COLLECTION.find_one({"room_id": room_id})
        if not room:
            return jsonify({
                "code": 404,
                "message": "房间不存在"
            }), 404
        
        # 检查房间是否已满
        if room["current_participants"] >= room["max_participants"]:
            return jsonify({
                "code": 409,
                "message": "房间已满"
            }), 409
        
        # 检查用户是否已在房间中
        existing_session = USER_SESSIONS_COLLECTION.find_one({
            "user_id": user_id,
            "room_id": room_id
        })
        
        if existing_session:
            return jsonify({
                "code": 409,
                "message": "您已在此房间中"
            }), 409
        
        # 创建用户会话
        session_data = {
            "user_id": user_id,
            "room_id": room_id,
            "session_id": room["session_id"],
            "role": role,
            "country_id": country_id if role == "participant" else None,
            "joined_at": datetime.now(UTC).isoformat() + "Z",
            "last_active": datetime.now(UTC).isoformat() + "Z",
            "status": "active"
        }
        
        USER_SESSIONS_COLLECTION.insert_one(session_data)
        
        # 更新房间参与人数
        MEETING_ROOMS_COLLECTION.update_one(
            {"room_id": room_id},
            {
                "$inc": {"current_participants": 1},
                "$push": {"participants": {
                    "user_id": user_id,
                    "role": role,
                    "country_id": country_id,
                    "joined_at": session_data["joined_at"]
                }}
            }
        )
        
        return jsonify({
            "code": 200,
            "message": "成功加入房间",
            "data": {
                "room_id": room_id,
                "session_id": room["session_id"],
                "role": role,
                "country_id": country_id
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"加入房间失败: {str(e)}"
        }), 500

# =========================
# 页面路由
# =========================

@app.route('/meeting-hall')
def meeting_hall():
    """会议大厅 - 多用户入口"""
    return render_template('meeting_hall.html')

@app.route('/landing')
def landing_page():
    """初始入口：身份选择 + 国家选择 + 快捷跳转"""
    # 重定向到新的系统主页
    return redirect('/system-home')

@app.route('/country-select')
def country_select_page():
    """国家选择页面 - 重定向到新的代表门户"""
    return redirect('/country-portal')

@app.route('/chairman-selection')
def chairman_selection_page():
    """主席选择页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    return render_template('chairman_selection.html', 
                         session_id=session_id, 
                         committee_name=committee_name, 
                         agenda=agenda)

@app.route('/voting-mechanism')
def voting_mechanism_page():
    """投票机制选择页面（临时跳过，直接进入主席点名）"""
    session_id = request.args.get("session_id", "")
    # 仍保留从入口可能传入的参数，便于后续扩展
    _ = request.args.get("chairman_id", "")
    _ = request.args.get("chairman_name", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    return redirect(f"/chairman-rollcall?session_id={session_id}&committee_name={committee_name}&agenda={agenda}")

@app.route('/real-time-voting')
def real_time_voting_page():
    """实时投票显示页面"""
    session_id = request.args.get("session_id", "")
    chairman_id = request.args.get("chairman_id", "")
    chairman_name = request.args.get("chairman_name", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    mechanism_name = request.args.get("mechanism_name", "")
    mechanism_requirement = request.args.get("mechanism_requirement", "")
    return render_template('real_time_voting.html',
                         session_id=session_id,
                         chairman_id=chairman_id,
                         chairman_name=chairman_name,
                         committee_name=committee_name,
                         agenda=agenda,
                         mechanism_name=mechanism_name,
                         mechanism_requirement=mechanism_requirement)

@app.route('/chairman-rollcall')
def chairman_rollcall_page():
    """主席点名页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    mechanism_name = request.args.get("mechanism_name", "")
    mechanism_requirement = request.args.get("mechanism_requirement", "")
    return render_template('chairman_rollcall.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         mechanism_name=mechanism_name,
                         mechanism_requirement=mechanism_requirement)

@app.route('/chairman-file-submission')
def chairman_file_submission_page():
    """主席文件提交监控页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    mechanism_name = request.args.get("mechanism_name", "")
    mechanism_requirement = request.args.get("mechanism_requirement", "")
    return render_template('chairman_file_submission.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         mechanism_name=mechanism_name,
                         mechanism_requirement=mechanism_requirement)

# =========================
# 与会国代表门户
# =========================

@app.route('/country-portal')
def country_portal():
    """与会国代表门户"""
    return send_from_directory('.', 'country_portal.html')


@app.route('/chairman-vote-monitoring')
def chairman_vote_monitoring_page():
    """主席投票监控页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    mechanism_name = request.args.get("mechanism_name", "协商一致")
    mechanism_requirement = request.args.get("mechanism_requirement", "要求：100% 同意")
    return render_template('chairman_vote_monitoring.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         mechanism_name=mechanism_name,
                         mechanism_requirement=mechanism_requirement)

@app.route('/chairman-declaration')
def chairman_declaration_page():
    """主席共同宣言管理页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    return render_template('chairman_declaration.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda)

# =========================
# 投票监控相关API
# =========================

@app.route('/api/send_vote_reminder', methods=['POST'])
def api_send_vote_reminder():
    """发送投票提醒"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        
        # 这里可以集成实际的提醒发送逻辑
        return jsonify({
            "code": 200,
            "message": "投票提醒已发送"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"发送提醒失败: {str(e)}"
        }), 500

@app.route('/api/extend_vote_deadline', methods=['POST'])
def api_extend_vote_deadline():
    """延长投票截止时间"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        new_deadline = data.get("new_deadline")
        
        if not new_deadline:
            return jsonify({
                "code": 400,
                "message": "请提供新的截止时间"
            }), 400
        
        cols = get_cols_by_session(session_id)
        cols["settings"].update_one(
            {"session_id": session_id},
            {"$set": {"vote_deadline": new_deadline}},
            upsert=True
        )
        
        return jsonify({
            "code": 200,
            "message": "投票截止时间已更新"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"更新截止时间失败: {str(e)}"
        }), 500

@app.route('/api/force_end_voting', methods=['POST'])
def api_force_end_voting():
    """强制结束投票并直接完成投票流程，进入共同宣言阶段"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")

        cols = get_cols_by_session(session_id)

        print(f"\n{'='*60}")
        print(f"🔧 开始强制结束投票并完成流程，session_id: {session_id}")
        print(f"{'='*60}")

        # 步骤1: 将所有未完成的投票标记为弃权
        print(f"\n📋 步骤1: 标记未完成的投票为弃权")
        uncompleted_count = cols["db"]["file_vote_details"].count_documents(
            {"session_id": session_id, "vote_result": {"$exists": False}}
        )

        result = cols["db"]["file_vote_details"].update_many(
            {"session_id": session_id, "vote_result": {"$exists": False}},
            {"$set": {
                "vote_result": "abstain",
                "voted_at": datetime.now(UTC).isoformat() + "Z",
                "forced": True,
                "forced_end": True
            }}
        )

        print(f"✅ 已标记 {result.modified_count} 个未完成投票为弃权")

        # 步骤2: 构建投票矩阵
        print(f"\n📋 步骤2: 构建投票矩阵")
        vote_matrix = {}
        file_vote_details = list(cols["db"]["file_vote_details"].find({"session_id": session_id}))

        for vote in file_vote_details:
            country_id = vote.get("country_id")
            file_id = vote.get("file_id")
            vote_result = vote.get("vote_result")

            if country_id and file_id and vote_result:
                if country_id not in vote_matrix:
                    vote_matrix[country_id] = {}
                vote_matrix[country_id][file_id] = vote_result

        print(f"📊 投票矩阵构建完成，包含 {len(vote_matrix)} 个国家，{len(file_vote_details)} 个投票记录")

        # 步骤3: 计算每个文件的投票结果
        print(f"\n📋 步骤3: 计算文件投票结果")
        file_results = {}
        for vote in file_vote_details:
            file_id = vote.get("file_id")
            vote_result = vote.get("vote_result")

            if file_id and vote_result:
                if file_id not in file_results:
                    file_results[file_id] = {'agree': 0, 'disagree': 0, 'abstain': 0}
                file_results[file_id][vote_result] += 1

        print(f"📈 计算完成，共处理 {len(file_results)} 个文件的投票结果")

        # 步骤4: 保存投票完成记录
        print(f"\n📋 步骤4: 保存投票完成记录")
        completed_at = datetime.now(UTC).isoformat() + "Z"
        voting_record = {
            "session_id": session_id,
            "vote_matrix": vote_matrix,
            "file_results": file_results,
            "completed_at": completed_at,
            "status": "completed",
            "force_ended": True,
            "uncompleted_count": uncompleted_count
        }

        cols["db"]["voting_records"].insert_one(voting_record)
        print(f"💾 投票记录已保存到 voting_records 集合")

        # 步骤5: 处理通过的文件，保存到passed_files集合和submissions集合
        print(f"\n📋 步骤5: 处理通过的文件")
        passed_files_list = []
        for file_id, results in file_results.items():
            # 判断是否通过（同意票 > 反对票）
            is_passed = results['agree'] > results['disagree']

            if is_passed:
                print(f"🔍 处理通过的文件: file_id={file_id}")

                # 多种方式获取文件信息
                file_info = None
                country_id = None
                file_name = None
                original_name = None

                # 方法1: 从temp_files获取
                temp_file = cols["db"]["temp_files"].find_one({"file_id": file_id, "session_id": session_id})
                if temp_file:
                    print(f"  ✅ 从temp_files找到文件信息")
                    file_info = temp_file
                    file_name = temp_file.get("saved_name") or temp_file.get("file_name", "")
                    original_name = temp_file.get("original_name", file_name)
                    country_id = temp_file.get("country_id", "")

                # 方法2: 从vote_files获取
                if not file_info:
                    vote_file = cols["db"]["vote_files"].find_one({"file_id": file_id, "session_id": session_id})
                    if vote_file:
                        print(f"  ✅ 从vote_files找到文件信息")
                        file_info = vote_file
                        file_name = vote_file.get("saved_name") or vote_file.get("file_name", "")
                        original_name = vote_file.get("original_name", file_name)
                        country_id = vote_file.get("country_id", "")

                # 方法3: 从file_vote_details反查country_id，再从submissions获取
                if not file_info:
                    print(f"  ⚠️  temp_files和vote_files都没找到，尝试从file_vote_details反查...")
                    vote_detail = cols["db"]["file_vote_details"].find_one({
                        "file_id": file_id,
                        "session_id": session_id
                    })
                    if vote_detail:
                        country_id = vote_detail.get("country_id", "")
                        print(f"  🔍 从file_vote_details找到country_id: {country_id}")

                        # 从submissions获取文件信息
                        if country_id:
                            submission = cols["submissions"].find_one({
                                "session_id": session_id,
                                "country_id": country_id
                            })
                            if submission:
                                print(f"  ✅ 从submissions找到文件信息")
                                file_name = submission.get("file_name", "")
                                original_name = file_name
                                file_info = submission

                # 如果找到了文件信息，保存到passed_files
                if file_info and file_name:
                    print(f"  📄 文件名: {file_name}")
                    print(f"  🌍 国家: {country_id}")

                    # 保存到passed_files集合（专门用于共同宣言生成）
                    passed_file_record = {
                        "session_id": session_id,
                        "file_id": file_id,
                        "file_name": file_name,
                        "original_name": original_name,
                        "country_id": country_id,
                        "vote_agree": results['agree'],
                        "vote_disagree": results['disagree'],
                        "vote_abstain": results['abstain'],
                        "passed_at": completed_at,
                        "status": "passed",
                        "force_passed": True
                    }

                    cols["db"]["passed_files"].update_one(
                        {"session_id": session_id, "file_id": file_id},
                        {"$set": passed_file_record},
                        upsert=True
                    )
                    print(f"  💾 已保存到passed_files集合")

                    # 同时更新submissions集合，标记为通过
                    if country_id:
                        cols["submissions"].update_one(
                            {"session_id": session_id, "country_id": country_id},
                            {"$set": {
                                "vote_passed": True,
                                "vote_status": "passed",
                                "vote_agree_count": results['agree'],
                                "vote_disagree_count": results['disagree'],
                                "vote_abstain_count": results['abstain'],
                                "vote_completed_at": completed_at,
                                "force_passed": True
                            }},
                            upsert=False
                        )
                        print(f"  💾 已更新submissions集合")

                    passed_files_list.append({
                        "file_id": file_id,
                        "file_name": file_name,
                        "original_name": original_name,
                        "country_id": country_id,
                        "vote_results": results
                    })

                    print(f"  ✅ 文件通过投票：{original_name} (file_id: {file_id}, 文件名: {file_name})")
                else:
                    print(f"  ❌ 警告：无法找到file_id={file_id}的文件信息！")
                    print(f"     同意票: {results['agree']}, 反对票: {results['disagree']}")

        print(f"📋 投票完成，共有 {len(passed_files_list)} 个文件通过（强制结束）")

        # 步骤6: 更新会议状态为宣言阶段
        print(f"\n📋 步骤6: 更新会议状态为宣言阶段")
        cols["settings"].update_one(
            {"session_id": session_id},
            {"$set": {
                "meeting_phase": "declaration",
                "voting_completed": True,
                "voting_completed_at": completed_at,
                "force_ended_voting": True,
                "passed_files_count": len(passed_files_list)
            }},
            upsert=True
        )
        print(f"🎯 会议状态已更新为宣言阶段")

        print(f"\n{'='*60}")
        print(f"✅ 强制结束投票并完成流程成功！")
        print(f"📊 处理了 {len(file_results)} 个文件的投票")
        print(f"✅ {len(passed_files_list)} 个文件通过投票")
        print(f"⏭️  会议已进入共同宣言阶段")
        print(f"{'='*60}\n")

        return jsonify({
            "code": 200,
            "message": f"投票已强制结束并完成！{len(passed_files_list)}个文件通过，会议进入共同宣言阶段",
            "data": {
                "force_ended": True,
                "uncompleted_count": uncompleted_count,
                "passed_files_count": len(passed_files_list),
                "file_results": file_results,
                "passed_files": passed_files_list,
                "next_phase": "declaration"
            }
        })

    except Exception as e:
        print(f"❌ 强制结束投票并完成流程失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "code": 500,
            "message": f"强制结束投票失败: {str(e)}"
        }), 500

@app.route('/api/finalize_file_voting', methods=['POST'])
def api_finalize_file_voting():
    """完成文件投票流程"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        vote_matrix = data.get("vote_matrix", {})
        completed_at = data.get("completed_at")
        
        cols = get_cols_by_session(session_id)
        
        # 计算每个文件的投票结果
        file_results = {}
        for country_id, country_votes in vote_matrix.items():
            for file_id, vote_result in country_votes.items():
                if vote_result is not None:
                    if file_id not in file_results:
                        file_results[file_id] = {'agree': 0, 'disagree': 0, 'abstain': 0}
                    file_results[file_id][vote_result] += 1
        
        # 保存投票完成记录
        voting_record = {
            "session_id": session_id,
            "vote_matrix": vote_matrix,
            "file_results": file_results,
            "completed_at": completed_at,
            "status": "completed"
        }
        
        cols["db"]["voting_records"].insert_one(voting_record)
        
        # 【新增】处理通过的文件，保存到passed_files集合和submissions集合
        passed_files_list = []
        for file_id, results in file_results.items():
            # 判断是否通过（同意票 > 反对票）
            is_passed = results['agree'] > results['disagree']
            
            if is_passed:
                print(f"\n🔍 处理通过的文件: file_id={file_id}")
                
                # 【改进】多种方式获取文件信息
                file_info = None
                country_id = None
                file_name = None
                original_name = None
                
                # 方法1: 从temp_files获取
                temp_file = cols["db"]["temp_files"].find_one({"file_id": file_id, "session_id": session_id})
                if temp_file:
                    print(f"  ✅ 从temp_files找到文件信息")
                    file_info = temp_file
                    file_name = temp_file.get("saved_name") or temp_file.get("file_name", "")
                    original_name = temp_file.get("original_name", file_name)
                    country_id = temp_file.get("country_id", "")
                
                # 方法2: 从vote_files获取
                if not file_info:
                    vote_file = cols["db"]["vote_files"].find_one({"file_id": file_id, "session_id": session_id})
                    if vote_file:
                        print(f"  ✅ 从vote_files找到文件信息")
                        file_info = vote_file
                        file_name = vote_file.get("saved_name") or vote_file.get("file_name", "")
                        original_name = vote_file.get("original_name", file_name)
                        country_id = vote_file.get("country_id", "")
                
                # 方法3: 从file_vote_details反查country_id，再从submissions获取
                if not file_info:
                    print(f"  ⚠️  temp_files和vote_files都没找到，尝试从file_vote_details反查...")
                    vote_detail = cols["db"]["file_vote_details"].find_one({
                        "file_id": file_id, 
                        "session_id": session_id
                    })
                    if vote_detail:
                        country_id = vote_detail.get("country_id", "")
                        print(f"  🔍 从file_vote_details找到country_id: {country_id}")
                        
                        # 从submissions获取文件信息
                        if country_id:
                            submission = cols["submissions"].find_one({
                                "session_id": session_id,
                                "country_id": country_id
                            })
                            if submission:
                                print(f"  ✅ 从submissions找到文件信息")
                                file_name = submission.get("file_name", "")
                                original_name = file_name
                                file_info = submission
                
                # 方法4: 遍历所有submissions，找到匹配的文件
                if not file_info:
                    print(f"  ⚠️  尝试遍历submissions查找...")
                    all_submissions = cols["submissions"].find({"session_id": session_id})
                    for sub in all_submissions:
                        # 如果文件名中包含file_id，或者其他匹配逻辑
                        if sub.get("file_name") and file_id in str(sub.get("_id", "")):
                            print(f"  ✅ 在submissions中找到匹配文件")
                            file_name = sub.get("file_name", "")
                            original_name = file_name
                            country_id = sub.get("country_id", "")
                            file_info = sub
                            break
                
                # 如果找到了文件信息，保存到passed_files
                if file_info and file_name:
                    print(f"  📄 文件名: {file_name}")
                    print(f"  🌍 国家: {country_id}")
                    
                    # 保存到passed_files集合（专门用于共同宣言生成）
                    passed_file_record = {
                        "session_id": session_id,
                        "file_id": file_id,
                        "file_name": file_name,  # 存储文件名字符串
                        "original_name": original_name,
                        "country_id": country_id,
                        "vote_agree": results['agree'],
                        "vote_disagree": results['disagree'],
                        "vote_abstain": results['abstain'],
                        "passed_at": completed_at,
                        "status": "passed"
                    }
                    
                    cols["db"]["passed_files"].update_one(
                        {"session_id": session_id, "file_id": file_id},
                        {"$set": passed_file_record},
                        upsert=True
                    )
                    print(f"  💾 已保存到passed_files集合")
                    
                    # 同时更新submissions集合，标记为通过
                    if country_id:
                        cols["submissions"].update_one(
                            {"session_id": session_id, "country_id": country_id},
                            {"$set": {
                                "vote_passed": True,
                                "vote_status": "passed",
                                "vote_agree_count": results['agree'],
                                "vote_disagree_count": results['disagree'],
                                "vote_abstain_count": results['abstain'],
                                "vote_completed_at": completed_at
                            }},
                            upsert=False
                        )
                        print(f"  💾 已更新submissions集合")
                    
                    passed_files_list.append({
                        "file_id": file_id,
                        "file_name": file_name,
                        "original_name": original_name,
                        "country_id": country_id
                    })
                    
                    print(f"  ✅ 文件通过投票：{original_name} (file_id: {file_id}, 文件名: {file_name})")
                else:
                    print(f"  ❌ 警告：无法找到file_id={file_id}的文件信息！")
                    print(f"     同意票: {results['agree']}, 反对票: {results['disagree']}")
                    print(f"     请检查temp_files、vote_files或submissions集合中是否有此文件")
        
        print(f"📋 投票完成，共有 {len(passed_files_list)} 个文件通过")
        
        return jsonify({
            "code": 200,
            "message": "文件投票流程已完成",
            "data": {
                "file_results": file_results,
                "passed_files": passed_files_list,
                "passed_count": len(passed_files_list)
            }
        })
        
    except Exception as e:
        print(f"❌ 完成投票失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "code": 500,
            "message": f"完成投票失败: {str(e)}"
        }), 500

@app.route('/api/rebuild_passed_files', methods=['POST'])
def api_rebuild_passed_files():
    """🔧 补救API：基于file_vote_details重建passed_files集合"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        
        print(f"\n{'='*60}")
        print(f"🔧 开始重建passed_files集合，session_id: {session_id}")
        print(f"{'='*60}")
        
        cols = get_cols_by_session(session_id)
        
        # 1. 从file_vote_details获取所有投票记录
        vote_details = list(cols["db"]["file_vote_details"].find({"session_id": session_id}))
        print(f"\n📊 找到 {len(vote_details)} 条投票记录")
        
        if not vote_details:
            return jsonify({
                "code": 400,
                "message": "没有找到投票记录"
            }), 400
        
        # 2. 统计每个文件的投票结果
        file_results = {}
        for vote in vote_details:
            file_id = vote.get("file_id")
            vote_result = vote.get("vote_result")
            
            if file_id and vote_result:
                if file_id not in file_results:
                    file_results[file_id] = {'agree': 0, 'disagree': 0, 'abstain': 0, 'country_id': vote.get("country_id")}
                file_results[file_id][vote_result] += 1
        
        print(f"\n📈 投票统计：")
        for file_id, results in file_results.items():
            print(f"  {file_id}: 同意={results['agree']}, 反对={results['disagree']}, 弃权={results['abstain']}")
        
        # 3. 判断哪些文件通过，并保存到passed_files
        passed_count = 0
        passed_files_list = []
        
        for file_id, results in file_results.items():
            is_passed = results['agree'] > results['disagree']
            
            if is_passed:
                print(f"\n✅ 文件 {file_id} 通过投票")
                
                # 获取country_id
                country_id = results.get('country_id', '')
                
                # 从submissions获取文件信息
                submission = None
                if country_id:
                    submission = cols["submissions"].find_one({
                        "session_id": session_id,
                        "country_id": country_id
                    })
                
                # 如果没找到，尝试从vote_files获取
                if not submission:
                    vote_file = cols["db"]["vote_files"].find_one({
                        "file_id": file_id,
                        "session_id": session_id
                    })
                    if vote_file:
                        country_id = vote_file.get("country_id", "")
                        submission = cols["submissions"].find_one({
                            "session_id": session_id,
                            "country_id": country_id
                        })
                
                # 如果还没找到，尝试从temp_files获取
                if not submission:
                    temp_file = cols["db"]["temp_files"].find_one({
                        "file_id": file_id,
                        "session_id": session_id
                    })
                    if temp_file:
                        country_id = temp_file.get("country_id", "")
                        submission = cols["submissions"].find_one({
                            "session_id": session_id,
                            "country_id": country_id
                        })
                
                if submission:
                    file_name = submission.get("file_name", "")
                    original_name = file_name
                    
                    print(f"  📄 文件名: {file_name}")
                    print(f"  🌍 国家: {country_id}")
                    
                    # 保存到passed_files集合
                    passed_file_record = {
                        "session_id": session_id,
                        "file_id": file_id,
                        "file_name": file_name,
                        "original_name": original_name,
                        "country_id": country_id,
                        "vote_agree": results['agree'],
                        "vote_disagree": results['disagree'],
                        "vote_abstain": results['abstain'],
                        "passed_at": datetime.now(UTC).isoformat() + "Z",
                        "status": "passed"
                    }
                    
                    cols["db"]["passed_files"].update_one(
                        {"session_id": session_id, "file_id": file_id},
                        {"$set": passed_file_record},
                        upsert=True
                    )
                    
                    # 更新submissions集合
                    cols["submissions"].update_one(
                        {"session_id": session_id, "country_id": country_id},
                        {"$set": {
                            "vote_passed": True,
                            "vote_status": "passed",
                            "vote_agree_count": results['agree'],
                            "vote_disagree_count": results['disagree'],
                            "vote_abstain_count": results['abstain'],
                            "vote_completed_at": datetime.now(UTC).isoformat() + "Z"
                        }},
                        upsert=False
                    )
                    
                    passed_count += 1
                    passed_files_list.append({
                        "file_id": file_id,
                        "file_name": file_name,
                        "country_id": country_id
                    })
                    
                    print(f"  💾 已保存到passed_files")
                else:
                    print(f"  ⚠️  警告：找不到对应的submission记录")
        
        print(f"\n{'='*60}")
        print(f"✅ 重建完成！共有 {passed_count} 个文件通过")
        print(f"{'='*60}")
        
        return jsonify({
            "code": 200,
            "message": f"成功重建passed_files集合，共有 {passed_count} 个文件通过",
            "data": {
                "passed_count": passed_count,
                "passed_files": passed_files_list
            }
        })
        
    except Exception as e:
        print(f"❌ 重建passed_files失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "code": 500,
            "message": f"重建失败: {str(e)}"
        }), 500

@app.route('/api/create_new_session', methods=['POST'])
def api_create_new_session():
    """创建新会议"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "")
        committee_name = data.get("committee_name", "")
        agenda = data.get("agenda", "")
        created_by = data.get("created_by", "")
        
        if not session_id:
            return jsonify({
                "code": 400,
                "message": "会议编号不能为空"
            }), 400
        
        cols = get_cols_by_session(session_id)
        
        # 检查会议是否已存在
        existing_session = cols["settings"].find_one({"session_id": session_id})
        if existing_session:
            return jsonify({
                "code": 409,
                "message": "会议编号已存在"
            }), 409
        
        # 创建新会议记录
        session_record = {
            "session_id": session_id,
            "committee_name": committee_name,
            "agenda": agenda,
            "created_by": created_by,
            "created_at": datetime.now(UTC).isoformat() + "Z",
            "status": "active",
            "participants": [],  # 初始为空，后续添加
            "settings": {
                "voting_mechanism": None,
                "chairman_id": None,
                "chairman_name": None
            },
            # 新增：会议状态管理
            "meeting_state": {
                "current_phase": "init",
                "phase_history": [
                    {
                        "phase": "init",
                        "started_at": datetime.now(UTC).isoformat() + "Z",
                        "completed_at": None
                    }
                ],
                "phase_locks": {
                    "rollcall": False,
                    "file_submission": False,
                    "motion": False,
                    "voting": False,
                    "declaration": False
                }
            },
            # 主席控制权限
            "chairman_controls": {
                "can_advance_phase": True,
                "can_go_back": False,
                "can_pause_meeting": True,
                "can_modify_participants": True
            }
        }
        
        # 保存到数据库（会期专属库）
        cols["settings"].insert_one(session_record)

        # 同步写入全局会议设置集合，供跨设备加入和国家保存使用
        try:
            # 若全局集合中已存在则不重复写入
            existing_global = col_settings.find_one({"session_id": session_id})
            if not existing_global:
                col_settings.insert_one({
                    "session_id": session_id,
                    "committee_name": committee_name,
                    "agenda": agenda,
                    "created_by": created_by,
                    "created_at": datetime.now(UTC).isoformat() + "Z",
                    "status": "active",
                    "participants": [],
                    "meeting_state": session_record.get("meeting_state", {}),
                    "chairman_controls": session_record.get("chairman_controls", {}),
                })
            else:
                # 轻量更新基本信息，避免被旧数据覆盖
                col_settings.update_one(
                    {"session_id": session_id},
                    {"$set": {
                        "committee_name": committee_name,
                        "agenda": agenda,
                        "status": "active"
                    }}
                )
        except Exception as e:
            # 不阻断会议创建流程，但记录日志便于排查
            print(f"同步写入全局会议设置失败: {e}")
        
        # 如果是五位数会议编号，确保创建独立数据库的索引
        if session_id and len(session_id) == 5 and session_id.isdigit():
            ensure_indexes(cols["db"])
        
        return jsonify({
            "code": 200,
            "message": f"新会议 {session_id} 创建成功",
            "data": {
                "session_id": session_id,
                "committee_name": committee_name,
                "agenda": agenda
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"创建新会议失败: {str(e)}"
        }), 500

# =========================
# 共同宣言管理相关API
# =========================

@app.route('/api/finalize_declaration', methods=['POST'])
def api_finalize_declaration():
    """确认定稿共同宣言"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        declaration = data.get("declaration", "")
        
        if not declaration.strip():
            return jsonify({
                "code": 400,
                "message": "宣言内容不能为空"
            }), 400
        
        cols = get_cols_by_session(session_id)
        
        # 更新宣言状态为定稿
        cols["declarations"].update_one(
            {"session_id": session_id},
            {"$set": {
                "declaration": declaration,
                "status": "finalized",
                "finalized_at": datetime.now(UTC).isoformat() + "Z"
            }},
            upsert=True
        )
        
        return jsonify({
            "code": 200,
            "message": "共同宣言已确认定稿"
        })

    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"定稿失败: {str(e)}"
        }), 500


@app.route('/api/start_file_submission', methods=['POST'])
def api_start_file_submission():
    """主席启动文件提交阶段，向与会国发送弹窗通知"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        committee_name = data.get("committee_name", "")
        agenda = data.get("agenda", "")

        cols = get_cols_by_session(session_id)
        if not cols:
            return jsonify({
                "code": 400,
                "message": "会议室不存在",
                "data": None
            })

        # 获取到场国家列表
        arrived_countries = []
        for col_name, col in cols.items():
            if col_name.startswith("country_"):
                country_data = col.get("country_data", {})
                if country_data.get("arrived", False):
                    arrived_countries.append({
                        "country_id": col_name,
                        "country_name": country_data.get("country_name", ""),
                        "country_flag": country_data.get("country_flag", "")
                    })

        if not arrived_countries:
            return jsonify({
                "code": 400,
                "message": "没有到场国家",
                "data": None
            })

        # 向到场国家发送文件提交通知
        notification_sent = 0
        for country in arrived_countries:
            try:
                # 向每个国家发送通知
                notification_data = {
                    "type": "file_submission_start",
                    "session_id": session_id,
                    "committee_name": committee_name,
                    "agenda": agenda,
                    "country_name": country["country_name"],
                    "message": f"主席已启动文件提交阶段，请尽快提交您的文件。"
                }

                # 查找该国家的列并发送通知
                country_col = cols.get(country["country_id"])
                if country_col:
                    # 发送通知到该国家的列
                    emit_to_room(country["country_id"], "notification", notification_data)
                    notification_sent += 1

            except Exception as e:
                print(f"向 {country['country_name']} 发送通知失败: {str(e)}")

        return jsonify({
            "code": 200,
            "message": f"文件提交阶段已启动，已向 {notification_sent} 个到场国家发送通知",
            "data": {
                "notified_countries": notification_sent,
                "total_arrived": len(arrived_countries)
            }
        })

    except Exception as e:
        print(f"启动文件提交阶段失败: {str(e)}")
        return jsonify({
            "code": 500,
            "message": f"启动文件提交阶段失败: {str(e)}",
            "data": None
        })


@app.route('/api/get_declarations', methods=['GET'])
def api_get_declarations():
    """获取共同宣言列表"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取该会期的所有宣言，按创建时间倒序
        declarations = list(cols["declarations"].find(
            {"session_id": session_id}
        ).sort("created_at", -1))
        
        # 转换ObjectId为字符串
        for declaration in declarations:
            declaration["_id"] = str(declaration["_id"])
        
        return jsonify({
            "code": 200,
            "message": "获取宣言列表成功",
            "data": declarations
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"获取宣言失败: {str(e)}"
        }), 500

# =========================
# 与会国页面路由
# =========================

@app.route('/country-motion')
def country_motion_page():
    """与会国动议参与页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    country_id = request.args.get("country_id", "")
    country_name = request.args.get("country_name", "")
    motion_country = request.args.get("motion_country", "")
    motion_text = request.args.get("motion_text", "")
    
    # 获取国家国旗 - 根据国家名称匹配国旗文件
    country_flag_url = "/static/flags/default.png"
    if country_name:
        try:
            # 国家名称到国旗文件的映射
            country_flag_map = {
                '中国': 'cn.png', '美国': 'us.png', '日本': 'jp.png', 
                '德国': 'de.png', '法国': 'fr.png', '英国': 'uk.png',
                '俄罗斯': 'ru.png', '印度': 'in.png', '巴西': 'br.png',
                '加拿大': 'ca.png', '澳大利亚': 'au.png', '韩国': 'kr.png',
                '意大利': 'it.png', '西班牙': 'es.png', '荷兰': 'nl.png',
                '瑞典': 'se.png', '挪威': 'no.png', '丹麦': 'dk.png',
                '芬兰': 'fi.png', '瑞士': 'ch.png', '奥地利': 'at.png',
                '比利时': 'be.png', '波兰': 'pl.png', '捷克': 'cz.png',
                '匈牙利': 'hu.png', '葡萄牙': 'pt.png', '希腊': 'gr.png',
                '土耳其': 'tr.png', '以色列': 'il.png', '沙特阿拉伯': 'sa.png',
                '阿联酋': 'ae.png', '埃及': 'eg.png', '南非': 'za.png',
                '尼日利亚': 'ng.png', '肯尼亚': 'ke.png', '摩洛哥': 'ma.png',
                '阿根廷': 'ar.png', '智利': 'cl.png', '墨西哥': 'mx.png',
                '哥伦比亚': 'co.png', '秘鲁': 'pe.png', '委内瑞拉': 've.png',
                '泰国': 'th.png', '马来西亚': 'my.png', '新加坡': 'sg.png',
                '印度尼西亚': 'id.png', '菲律宾': 'ph.png', '越南': 'vn.png',
                '新西兰': 'nz.png', '乌拉圭': 'uy.png'
            }
            
            flag_filename = country_flag_map.get(country_name)
            if flag_filename:
                flag_path = os.path.join('app', 'static', 'flags', flag_filename)
                if os.path.exists(flag_path):
                    country_flag_url = f"/static/flags/{flag_filename}"
        except:
            pass  # 如果查询失败，使用默认国旗
    
    return render_template('country_motion.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         country_id=country_id,
                         country_name=country_name,
                         country_flag_url=country_flag_url,
                         motion_country=motion_country,
                         motion_text=motion_text)

@app.route('/country-file-vote')
def country_file_vote_page():
    """与会国文件投票页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    country_id = request.args.get("country_id", "")
    country_name = request.args.get("country_name", "")
    mechanism_name = request.args.get("mechanism_name", "协商一致")
    mechanism_requirement = request.args.get("mechanism_requirement", "要求：100% 同意")
    
    # 获取国家国旗 - 根据国家名称匹配国旗文件
    country_flag_url = "/static/flags/default.png"
    if country_name:
        try:
            # 国家名称到国旗文件的映射
            country_flag_map = {
                '中国': 'cn.png', '美国': 'us.png', '日本': 'jp.png', 
                '德国': 'de.png', '法国': 'fr.png', '英国': 'uk.png',
                '俄罗斯': 'ru.png', '印度': 'in.png', '巴西': 'br.png',
                '加拿大': 'ca.png', '澳大利亚': 'au.png', '韩国': 'kr.png',
                '意大利': 'it.png', '西班牙': 'es.png', '荷兰': 'nl.png',
                '瑞典': 'se.png', '挪威': 'no.png', '丹麦': 'dk.png',
                '芬兰': 'fi.png', '瑞士': 'ch.png', '奥地利': 'at.png',
                '比利时': 'be.png', '波兰': 'pl.png', '捷克': 'cz.png',
                '匈牙利': 'hu.png', '葡萄牙': 'pt.png', '希腊': 'gr.png',
                '土耳其': 'tr.png', '以色列': 'il.png', '沙特阿拉伯': 'sa.png',
                '阿联酋': 'ae.png', '埃及': 'eg.png', '南非': 'za.png',
                '尼日利亚': 'ng.png', '肯尼亚': 'ke.png', '摩洛哥': 'ma.png',
                '阿根廷': 'ar.png', '智利': 'cl.png', '墨西哥': 'mx.png',
                '哥伦比亚': 'co.png', '秘鲁': 'pe.png', '委内瑞拉': 've.png',
                '泰国': 'th.png', '马来西亚': 'my.png', '新加坡': 'sg.png',
                '印度尼西亚': 'id.png', '菲律宾': 'ph.png', '越南': 'vn.png',
                '新西兰': 'nz.png', '乌拉圭': 'uy.png'
            }
            
            flag_filename = country_flag_map.get(country_name)
            if flag_filename:
                flag_path = os.path.join('app', 'static', 'flags', flag_filename)
                if os.path.exists(flag_path):
                    country_flag_url = f"/static/flags/{flag_filename}"
        except:
            pass  # 如果查询失败，使用默认国旗
    
    return render_template('country_file_vote.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         country_id=country_id,
                         country_name=country_name,
                         country_flag_url=country_flag_url,
                         mechanism_name=mechanism_name,
                         mechanism_requirement=mechanism_requirement)

@app.route('/country-declaration')
def country_declaration_page():
    """与会国共同宣言页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    country_id = request.args.get("country_id", "")
    country_name = request.args.get("country_name", "")
    declaration_time = request.args.get("declaration_time", "")
    
    # 获取国家国旗和emoji
    country_flag_url = "/static/flags/default.png"
    country_flag_emoji = "🏳️"
    if country_name:
        country_flag_map = {
            '德国': '/static/flags/de.png', '法国': '/static/flags/fr.png',
            '美国': '/static/flags/us.png', '中国': '/static/flags/cn.png',
            '日本': '/static/flags/jp.png', '英国': '/static/flags/uk.png'
        }
        country_emoji_map = {
            '德国': '🇩🇪', '法国': '🇫🇷', '美国': '🇺🇸', 
            '中国': '🇨🇳', '日本': '🇯🇵', '英国': '🇬🇧'
        }
        country_flag_url = country_flag_map.get(country_name, "/static/flags/default.png")
        country_flag_emoji = country_emoji_map.get(country_name, "🏳️")
    
    return render_template('country_declaration.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         country_id=country_id,
                         country_name=country_name,
                         country_flag_url=country_flag_url,
                         country_flag_emoji=country_flag_emoji,
                         declaration_time=declaration_time)

# =========================
# 与会国相关API
# =========================

@app.route('/api/request_speak', methods=['POST'])
def api_request_speak():
    """申请发言"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        country_name = data.get("country_name")
        
        # 这里可以实现实际的申请发言逻辑
        # 比如通知主席、加入发言队列等
        
        return jsonify({
            "code": 200,
            "message": f"{country_name}的发言申请已提交给主席"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"申请发言失败: {str(e)}"
        }), 500

@app.route('/api/request_time_extension', methods=['POST'])
def api_request_time_extension():
    """申请延时"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        country_name = data.get("country_name")
        reason = data.get("reason", "")
        
        # 保存延时申请记录
        cols = get_cols_by_session(session_id)
        extension_record = {
            "session_id": session_id,
            "country_id": country_id,
            "country_name": country_name,
            "reason": reason,
            "requested_at": datetime.now(UTC).isoformat() + "Z",
            "status": "pending"
        }
        
        cols["db"]["time_extension_requests"].insert_one(extension_record)
        
        return jsonify({
            "code": 200,
            "message": f"{country_name}的延时申请已提交给主席"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"申请延时失败: {str(e)}"
        }), 500

@app.route('/api/cast_file_vote', methods=['POST'])
def api_cast_file_vote():
    """投票给文件"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        file_id = data.get("file_id")
        vote_result = data.get("vote_result")
        
        if not all([country_id, file_id, vote_result]):
            return jsonify({
                "code": 400,
                "message": "缺少必要参数"
            }), 400
        
        if vote_result not in ['agree', 'disagree', 'abstain']:
            return jsonify({
                "code": 400,
                "message": "无效的投票选项"
            }), 400
        
        cols = get_cols_by_session(session_id)
        
        # 保存投票记录
        vote_record = {
            "session_id": session_id,
            "country_id": country_id,
            "file_id": file_id,
            "vote_result": vote_result,
            "voted_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        cols["db"]["file_vote_details"].update_one(
            {
                "session_id": session_id,
                "country_id": country_id,
                "file_id": file_id
            },
            {"$set": vote_record},
            upsert=True
        )
        
        return jsonify({
            "code": 200,
            "message": "投票已保存"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"投票失败: {str(e)}"
        }), 500

@app.route('/api/submit_country_votes', methods=['POST'])
def api_submit_country_votes():
    """提交国家的所有投票"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        votes = data.get("votes", {})

        cols = get_cols_by_session(session_id)

        # 同时保存到两个集合，确保主席监控页面能读取到数据
        current_time = datetime.now(UTC).isoformat() + "Z"

        # 1. 保存到 file_vote_details 集合（主席监控页面读取的数据源）
        for file_id, vote_result in votes.items():
            vote_record = {
                "session_id": session_id,
                "country_id": country_id,
                "file_id": file_id,
                "vote_result": vote_result,
                "voted_at": current_time
            }

            cols["db"]["file_vote_details"].update_one(
                {
                    "session_id": session_id,
                    "country_id": country_id,
                    "file_id": file_id
                },
                {"$set": vote_record},
                upsert=True
            )

        # 2. 保存到 country_vote_submissions 集合（用于记录提交状态）
        completion_record = {
            "session_id": session_id,
            "country_id": country_id,
            "votes": votes,
            "submitted_at": current_time,
            "status": "completed"
        }

        cols["db"]["country_vote_submissions"].update_one(
            {"session_id": session_id, "country_id": country_id},
            {"$set": completion_record},
            upsert=True
        )

        return jsonify({
            "code": 200,
            "message": "所有投票已提交"
        })

    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"提交投票失败: {str(e)}"
        }), 500

@app.route('/api/confirm_declaration', methods=['POST'])
def api_confirm_declaration():
    """确认共同宣言"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        country_name = data.get("country_name")
        
        cols = get_cols_by_session(session_id)
        
        # 保存确认记录
        confirmation_record = {
            "session_id": session_id,
            "country_id": country_id,
            "country_name": country_name,
            "confirmed_at": datetime.now(UTC).isoformat() + "Z",
            "status": "confirmed"
        }
        
        cols["db"]["declaration_confirmations"].update_one(
            {"session_id": session_id, "country_id": country_id},
            {"$set": confirmation_record},
            upsert=True
        )
        
        return jsonify({
            "code": 200,
            "message": f"{country_name}已确认共同宣言"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"确认失败: {str(e)}"
        }), 500

@app.route('/api/submit_declaration_feedback', methods=['POST'])
def api_submit_declaration_feedback():
    """提交宣言反馈"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        country_name = data.get("country_name")
        feedback = data.get("feedback", "")
        
        if not feedback.strip():
            return jsonify({
                "code": 400,
                "message": "反馈内容不能为空"
            }), 400
        
        cols = get_cols_by_session(session_id)
        
        # 保存反馈记录
        feedback_record = {
            "session_id": session_id,
            "country_id": country_id,
            "country_name": country_name,
            "feedback": feedback,
            "submitted_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        cols["db"]["declaration_feedback"].update_one(
            {"session_id": session_id, "country_id": country_id},
            {"$set": feedback_record},
            upsert=True
        )
        
        return jsonify({
            "code": 200,
            "message": f"{country_name}的反馈已提交"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"提交反馈失败: {str(e)}"
        }), 500

@app.route('/api/get_declaration_participation', methods=['GET'])
def api_get_declaration_participation():
    """获取国家的宣言参与情况"""
    try:
        session_id = request.args.get("session_id", "default")
        country_id = request.args.get("country_id", "")
        
        cols = get_cols_by_session(session_id)
        
        # 检查确认状态
        confirmation = cols["db"]["declaration_confirmations"].find_one({
            "session_id": session_id,
            "country_id": country_id
        })
        
        # 检查反馈状态
        feedback = cols["db"]["declaration_feedback"].find_one({
            "session_id": session_id,
            "country_id": country_id
        })
        
        return jsonify({
            "code": 200,
            "data": {
                "confirmed": confirmation is not None,
                "feedback_submitted": feedback is not None,
                "confirmed_at": confirmation.get("confirmed_at") if confirmation else None,
                "feedback_at": feedback.get("submitted_at") if feedback else None
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"获取参与情况失败: {str(e)}"
        }), 500

# =========================
# 文件提交监控相关API
# =========================

@app.route('/api/send_submission_reminder', methods=['POST'])
def api_send_submission_reminder():
    """发送文件提交提醒"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_ids = data.get("country_ids", [])  # 可以是单个或多个国家
        
        # 这里可以集成实际的提醒发送逻辑
        # 比如发送邮件、系统通知等
        
        if country_ids:
            return jsonify({
                "code": 200,
                "message": f"已向 {len(country_ids)} 个国家发送提醒",
                "data": {"reminded_countries": len(country_ids)}
            })
        else:
            return jsonify({
                "code": 400,
                "message": "未指定提醒国家"
            }), 400
            
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"发送提醒失败: {str(e)}"
        }), 500

@app.route('/api/extend_submission_deadline', methods=['POST'])
def api_extend_submission_deadline():
    """延长文件提交截止时间"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        new_deadline = data.get("new_deadline")
        
        if not new_deadline:
            return jsonify({
                "code": 400,
                "message": "请提供新的截止时间"
            }), 400
        
        # 保存新的截止时间到数据库
        cols = get_cols_by_session(session_id)
        cols["settings"].update_one(
            {"session_id": session_id},
            {"$set": {"submission_deadline": new_deadline}},
            upsert=True
        )
        
        return jsonify({
            "code": 200,
            "message": "截止时间已更新",
            "data": {"new_deadline": new_deadline}
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"更新截止时间失败: {str(e)}"
        }), 500

@app.route('/chairman-motion')
def chairman_motion_page():
    """主席动议管理页面"""
    session_id = request.args.get("session_id", "")
    committee_name = request.args.get("committee_name", "")
    agenda = request.args.get("agenda", "")
    motion_country = request.args.get("motion_country", "")
    return render_template('chairman_motion.html',
                         session_id=session_id,
                         committee_name=committee_name,
                         agenda=agenda,
                         motion_country=motion_country)

# =========================
# 动议管理相关API
# =========================

@app.route('/api/save_speaking_order', methods=['POST'])
def api_save_speaking_order():
    """保存发言顺序和实时状态"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        speaking_order = data.get("speaking_order", [])
        current_speaker = data.get("current_speaker", -1)
        is_timer_running = data.get("is_timer_running", False)
        current_timer = data.get("current_timer", 0)
        
        print(f"\n📥 [后端API] 收到保存请求:")
        print(f"  - session_id: {session_id}")
        print(f"  - current_speaker: {current_speaker}")
        print(f"  - is_timer_running: {is_timer_running}")
        print(f"  - current_timer: {current_timer}")
        print(f"  - speaking_order 数量: {len(speaking_order)}")
        
        cols = get_cols_by_session(session_id)
        
        # 保存发言顺序和计时器状态到数据库
        speaking_order_record = {
            "session_id": session_id,
            "speaking_order": speaking_order,
            "current_speaker": current_speaker,
            "is_timer_running": is_timer_running,
            "current_timer": current_timer,
            "updated_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        result = cols["db"]["speaking_orders"].update_one(
            {"session_id": session_id},
            {"$set": speaking_order_record},
            upsert=True
        )
        
        print(f"💾 [数据库] 保存结果: matched={result.matched_count}, modified={result.modified_count}, upserted_id={result.upserted_id}")
        
        return jsonify({
            "code": 200,
            "message": "发言顺序已保存"
        })
        
    except Exception as e:
        print(f"❌ [后端API] 保存失败: {str(e)}")
        return jsonify({
            "code": 500,
            "message": f"保存发言顺序失败: {str(e)}"
        }), 500

@app.route('/api/get_speaking_order', methods=['GET'])
def api_get_speaking_order():
    """获取发言顺序和实时状态"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        print(f"\n📤 [后端API] 收到获取请求: session_id={session_id}")
        
        speaking_order_doc = cols["db"]["speaking_orders"].find_one({"session_id": session_id})
        
        if speaking_order_doc:
            # 获取发言顺序数据
            speaking_order = speaking_order_doc.get("speaking_order", [])
            current_speaker = speaking_order_doc.get("current_speaker", -1)
            
            # 添加实时状态信息
            timer_state = speaking_order_doc.get("timer_state", {})
            is_timer_running = speaking_order_doc.get("is_timer_running", False)
            current_timer = speaking_order_doc.get("current_timer", 0)
            
            print(f"📦 [数据库] 读取到数据:")
            print(f"  - current_speaker: {current_speaker}")
            print(f"  - is_timer_running: {is_timer_running}")
            print(f"  - current_timer: {current_timer}")
            print(f"  - speaking_order 数量: {len(speaking_order)}")
            
            response_data = {
                "code": 200,
                "message": "获取发言顺序成功",
                "data": {
                    "speaking_order": speaking_order,
                    "current_speaker": current_speaker,
                    "is_timer_running": is_timer_running,
                    "current_timer": current_timer,
                    "timer_state": timer_state,
                    "updated_at": speaking_order_doc.get("updated_at", "")
                }
            }
            
            print(f"✅ [后端API] 返回数据成功")
            return jsonify(response_data)
        else:
            print(f"⚠️ [数据库] 未找到 session_id={session_id} 的数据")
            return jsonify({
                "code": 200,
                "message": "暂无发言顺序",
                "data": {
                    "speaking_order": [],
                    "current_speaker": -1,
                    "is_timer_running": False,
                    "current_timer": 0
                }
            })
        
    except Exception as e:
        print(f"❌ [后端API] 获取失败: {str(e)}")
        return jsonify({
            "code": 500,
            "message": f"获取发言顺序失败: {str(e)}"
        }), 500

@app.route('/api/complete_motion', methods=['POST'])
def api_complete_motion():
    """完成动议讨论"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        speaking_order = data.get("speaking_order", [])
        completed_at = data.get("completed_at")
        
        cols = get_cols_by_session(session_id)
        
        # 保存动议完成记录
        motion_record = {
            "session_id": session_id,
            "speaking_order": speaking_order,
            "completed_at": completed_at,
            "status": "completed",
            "total_speakers": len(speaking_order),
            "completed_speakers": len([s for s in speaking_order if s.get("status") == "completed"])
        }
        
        cols["db"]["motion_records"].insert_one(motion_record)
        
        return jsonify({
            "code": 200,
            "message": "动议讨论已完成"
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"完成动议失败: {str(e)}"
        }), 500

@app.route('/')
def index():
    """系统主页 - 新的系统入口"""
    return render_template('system_home.html')

@app.route('/system-home')
def system_home():
    """新系统主页"""
    return render_template('system_home.html')

@app.route('/country-portal2')
def country_portal2_page():
    """与会国门户页面"""
    return render_template('country_portal.html')

@app.route('/legacy', methods=['GET', 'POST'])
def legacy_index():
    """保留的旧版首页（已注释功能）"""
    # 如果是 POST 请求，获取表单数据并保存到 session
    if request.method == 'POST':
        session['committee_name'] = request.form.get('committee_name', '')
        session['agenda'] = request.form.get('agenda', '')

        return redirect(url_for('file_submit_page'))  # 提交后重定向到文件提交页面

    # 如果是 GET 请求，获取 session 中的数据，如果没有则使用默认值
    committee_name = session.get('committee_name', ' ')
    agenda = session.get('agenda', '')
    session_id = request.args.get("session_id", "default")

    return render_template('index.html', committee_name=committee_name, agenda=agenda, session_id=session_id)


@app.route('/file-submit')
def file_submit_page():
    """文件提交页面 - 重定向到新的文件提交页面"""
    return redirect('/file-upload-submit')

@app.route('/file-upload-submit')
def file_upload_submit_page():
    """新的文件提交页面"""
    session_id = request.args.get("session_id", "default")
    return render_template('file_upload_submit.html', session_id=session_id)

@app.route('/file-upload')
def file_upload_page():
    """文件上传管理页面 - 重定向到新的文件提交页面"""
    return redirect('/file-upload-submit')

@app.route('/rollcall')
def rollcall_page():
    """点名页面 - 重定向到主席点名页面"""
    return redirect('/chairman-rollcall')

@app.route('/declaration-generator')
def declaration_generator_page():
    """宣言生成测试页面 - 上传文件并生成宣言"""
    return render_template('declaration_generator.html')

@app.route('/motion')
@app.route('/country-motion')
def motion_page():
    """动议倒计时页面 / 与会国动议参与页面"""
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)
    country_id = request.args.get("country_id", "")
    country_name = request.args.get("country_name", "未知国家")

    # 获取会议设置
    sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", " ")
    agenda = sdoc.get("agenda", " ")

    # 获取该国家的提交内容作为动议内容
    motion_text = ""
    if country_id:
        submission = cols["submissions"].find_one({
            "session_id": session_id,
            "country_id": country_id
        })
        if submission:
            motion_text = submission.get("text", "")

    return render_template(
        'motion.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id,
        country_id=country_id,
        country_name=country_name,
        motion_text=motion_text
    )


@app.route('/vote')
@app.route('/country-file-vote')
def vote_page():
    """投票页面 / 与会国文件投票页面"""
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)
    country_id = request.args.get("country_id", "")
    country_name = request.args.get("country_name", "未知国家")

    # 获取会议设置
    sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", " ")
    agenda = sdoc.get("agenda", " ")

    return render_template(
        'vote.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id,
        country_id=country_id,
        country_name=country_name
    )

@app.route('/file-vote')
def file_vote_page():
    """文件投票页面"""
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)
    
    # 获取会议设置
    sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", "WTO委员会")
    agenda = sdoc.get("agenda", "贸易谈判")
    
    return render_template(
        'file_vote.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id
    )

@app.route('/comprehensive-vote')
def comprehensive_vote_page():
    """全面投票页面 - 每个国家对每个文件进行投票"""
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)
    
    # 获取会议设置
    sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", "WTO委员会")
    agenda = sdoc.get("agenda", "贸易谈判")
    
    return render_template(
        'comprehensive_vote.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id
    )

@app.route('/simple-vote')
def simple_vote_page():
    """简化投票页面 - 更简单可靠的投票界面"""
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)
    
    # 获取会议设置
    sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", "WTO委员会")
    agenda = sdoc.get("agenda", "贸易谈判")
    
    return render_template(
        'simple_vote.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id
    )

@app.route('/ultra-simple-vote')
def ultra_simple_vote_page():
    """超简单投票页面 - 最可靠的投票界面"""
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)
    
    # 获取会议设置
    sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", " ")
    agenda = sdoc.get("agenda", " ")
    
    return render_template(
        'ultra_simple_vote.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id
    )


# =========================
# API：国家列表（你原有的接口，保持不变）
# =========================
@app.route('/api/countries')
def get_countries():
    """
    获取国家列表
    - 如果 only_participants=true: 只返回已加入会议的国家(用于主席端显示参与国)
    - 否则: 返回所有国家(用于与会国选择国家)
    """
    session_id = request.args.get("session_id")
    only_participants = request.args.get("only_participants", "false").lower() == "true"
    
    cols = get_cols_by_session(session_id or "default")
    
    # 如果需要只返回参与国，则从 meeting_settings 获取已保存的参与国列表
    participants_data = []
    if only_participants and session_id:
        session_info = col_settings.find_one({"session_id": session_id})
        if session_info and session_info.get('participants'):
            # participants 是一个数组，包含 {country_id, country_name, country_flag, status}
            participants_data = session_info.get('participants', [])
            
            # 直接从 participants 构造返回数据
            data = []
            for p in participants_data:
                if p.get('status') == 'active':  # 只返回激活状态的参与国
                    data.append({
                        "id": p.get('country_id'),
                        "name": p.get('country_name', '未知国家'),
                        "flag_url": p.get('country_flag', '/static/flags/default.png')
                    })
            return jsonify({'code': 200, 'message': '获取参与国列表成功', 'data': data})
    
    # 如果不是只返回参与国，则返回所有国家（用于与会国门户选择）
    countries = list(cols["countries"].find())
    flag_dir = os.path.join(app.static_folder, "flags")

    data = []
    for country in countries:
        cid = str(country.get('_id'))
        name = country.get('country_name', '未知国家')

        # 兼容你原有字段：flag / code
        if country.get('flag'):
            filename = os.path.basename(country['flag'])
        elif country.get('code'):
            filename = f"{country['code'].lower()}.png"
        else:
            filename = None

        if filename and os.path.exists(os.path.join(flag_dir, filename)):
            flag_url = f"/static/flags/{filename}"
        else:
            flag_url = "/static/flags/default.png"

        data.append({"id": cid, "name": name, "flag_url": flag_url})

    return jsonify({'code': 200, 'message': '获取国家列表成功', 'data': data})


# =========================
# API：到场国家（点名）
# =========================
@app.route('/api/rollcall/arrived')
def api_rollcall_arrived():
    """
    获取当前会期到场国家：
    优先从 meeting_settings 的 participants 字段获取
    如果没有，则从 rollcall 集合获取（兼容旧数据）
    """
    session_id = request.args.get("session_id", "default")
    
    # 首先尝试从 meeting_settings 获取参与国
    session_info = col_settings.find_one({"session_id": session_id})
    if session_info and session_info.get('participants'):
        participants = session_info.get('participants', [])
        ids = [str(p['country_id']) for p in participants if p.get('status') == 'active']
        return jsonify({'code': 200, 'message': '获取到场国家成功', 'data': ids})
    
    # 如果没有参与国数据，则从 rollcall 集合获取（兼容旧数据）
    cols = get_cols_by_session(session_id)
    cur = cols["rollcall"].find({"session_id": session_id, "arrived": True}, {"country_id": 1})
    ids = []
    for doc in cur:
        cid = doc.get("country_id")
        # 统一转字符串，避免 ObjectId 带来前端对比问题
        if cid is not None:
            ids.append(str(cid))
    return jsonify({'code': 200, 'message': '获取到场国家成功', 'data': ids})

@app.route('/api/rollcall/update', methods=['POST'])
def api_rollcall_update():
    """更新点名状态"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        country_id = data.get("country_id")
        arrived = data.get("arrived", False)
        
        if not country_id:
            return jsonify({"code": 400, "message": "country_id 不能为空"}), 400
        
        cols = get_cols_by_session(session_id)
        
        # 更新或插入点名记录
        result = cols["rollcall"].update_one(
            {"session_id": session_id, "country_id": country_id},
            {"$set": {
                "session_id": session_id,
                "country_id": country_id,
                "arrived": arrived, 
                "updated_at": datetime.now(UTC).isoformat() + "Z"
            }},
            upsert=True
        )
        
        return jsonify({"code": 200, "message": "点名状态更新成功"})
        
    except Exception as e:
        print(f"更新点名状态时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"更新失败: {str(e)}"}), 500

@app.route('/api/rollcall/batch_update', methods=['POST'])
def api_rollcall_batch_update():
    """批量更新点名状态"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        updates = data.get("updates", [])
        
        if not updates:
            return jsonify({"code": 400, "message": "updates 不能为空"}), 400
        
        cols = get_cols_by_session(session_id)
        
        # 批量更新
        updated_count = 0
        for update in updates:
            country_id = update.get("country_id")
            arrived = update.get("arrived", False)
            
            if country_id:
                result = cols["rollcall"].update_one(
                    {"session_id": session_id, "country_id": country_id},
                    {"$set": {
                        "session_id": session_id,
                        "country_id": country_id,
                        "arrived": arrived, 
                        "updated_at": datetime.now(UTC).isoformat() + "Z"
                    }},
                    upsert=True
                )
                updated_count += 1
        return jsonify({"code": 200, "message": f"批量更新点名状态成功，共更新{updated_count}条记录"})
        
    except Exception as e:
        print(f"批量更新点名状态时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"批量更新失败: {str(e)}"}), 500

@app.route('/api/rollcall/statistics')
def api_rollcall_statistics():
    """获取点名统计信息"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取应出席国家总数
        sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
        total_countries = len(sdoc.get("participants", []))
        
        # 获取已出席国家数量
        arrived_count = cols["rollcall"].count_documents({"session_id": session_id, "arrived": True})
        
        # 获取未出席国家数量
        absent_count = cols["rollcall"].count_documents({"session_id": session_id, "arrived": False})
        
        # 计算出席率
        attendance_rate = (arrived_count / total_countries * 100) if total_countries > 0 else 0
        
        return jsonify({
            "code": 200,
            "data": {
                "total_countries": total_countries,
                "arrived_count": arrived_count,
                "absent_count": absent_count,
                "attendance_rate": round(attendance_rate, 1)
            }
        })
        
    except Exception as e:
        print(f"获取点名统计时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取统计失败: {str(e)}"}), 500

@app.route('/api/save_meeting_settings', methods=['POST'])
def api_save_meeting_settings():
    """保存会议设置"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        committee_name = data.get("committee_name", "").strip()
        agenda = data.get("agenda", "").strip()
        
        if not committee_name or not agenda:
            return jsonify({"code": 400, "message": "会议名称和议题不能为空"}), 400
        
        # 更新或插入会议设置
        cols["settings"].update_one(
            {"session_id": session_id},
            {
                "$set": {
                    "committee_name": committee_name,
                    "agenda": agenda,
                    "updated_at": datetime.now(UTC)
                }
            },
            upsert=True
        )
        
        return jsonify({"code": 200, "message": "会议设置保存成功"})
        
    except Exception as e:
        print(f"保存会议设置时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/save_participants', methods=['POST'])
def api_save_participants():
    """保存应出席国家"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        participants = data.get("participants", [])
        
        if not participants:
            return jsonify({"code": 400, "message": "应出席国家不能为空"}), 400
        
        # 更新或插入会议设置
        cols["settings"].update_one(
            {"session_id": session_id},
            {
                "$set": {
                    "participants": participants,
                    "updated_at": datetime.now(UTC)
                }
            },
            upsert=True
        )
        
        return jsonify({"code": 200, "message": "应出席国家保存成功"})
        
    except Exception as e:
        print(f"保存应出席国家时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/get_session_info')
def api_get_session_info():
    """获取会议信息"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取会议设置
        settings = cols["settings"].find_one({"session_id": session_id})
        
        if settings:
            # 获取会议状态，默认为active
            meeting_status = "active"

            # 检查是否有会议状态记录
            status_record = cols["db"]["meeting_status"].find_one({"session_id": session_id})
            if status_record:
                meeting_status = status_record.get("status", "active")

            return jsonify({
                "code": 200,
                "message": "获取会议信息成功",
                "data": {
                    "committee_name": settings.get("committee_name", ""),
                    "agenda": settings.get("agenda", ""),
                    "participants": settings.get("participants", []),
                    "status": meeting_status,
                    "created_at": settings.get("created_at"),
                    "session_id": session_id
                }
            })
        else:
            return jsonify({
                "code": 404,
                "message": "会议信息不存在"
            })
            
    except Exception as e:
        return jsonify({"code": 500, "message": f"获取会议信息失败: {str(e)}"}), 500

@app.route('/api/save_chairman', methods=['POST'])
def api_save_chairman():
    """保存主席选择"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        chairman_id = data.get("chairman_id")
        chairman_name = data.get("chairman_name")
        committee_name = data.get("committee_name", "")
        agenda = data.get("agenda", "")
        
        if not chairman_id or not chairman_name:
            return jsonify({"code": 400, "message": "主席信息不完整"}), 400
        
        # 获取数据库集合
        cols = get_cols_by_session(session_id)
        
        # 保存主席信息到会议设置
        chairman_doc = {
            "session_id": session_id,
            "chairman_id": chairman_id,  # 直接保存字符串，不转换为ObjectId
            "chairman_name": chairman_name,
            "committee_name": committee_name,
            "agenda": agenda,
            "updated_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        # 更新或插入设置
        cols["settings"].update_one(
            {"session_id": session_id},
            {"$set": chairman_doc},
            upsert=True
        )
        
        return jsonify({"code": 200, "message": "主席选择保存成功"})
        
    except Exception as e:
        return jsonify({"code": 500, "message": f"保存主席选择失败: {str(e)}"}), 500

@app.route('/api/save_voting_mechanism', methods=['POST'])
def api_save_voting_mechanism():
    """保存投票机制"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        chairman_id = data.get("chairman_id")
        chairman_name = data.get("chairman_name")
        committee_name = data.get("committee_name", "")
        agenda = data.get("agenda", "")
        mechanism_type = data.get("mechanism_type")
        mechanism_name = data.get("mechanism_name")
        required_percentage = data.get("required_percentage")
        mechanism_description = data.get("mechanism_description", "")
        
        if not mechanism_type or not mechanism_name or required_percentage is None:
            return jsonify({"code": 400, "message": "投票机制信息不完整"}), 400
        
        # 获取数据库集合
        cols = get_cols_by_session(session_id)
        
        # 保存投票机制信息
        mechanism_doc = {
            "session_id": session_id,
            "chairman_id": chairman_id,  # 直接保存字符串，不转换为ObjectId
            "chairman_name": chairman_name,
            "committee_name": committee_name,
            "agenda": agenda,
            "mechanism_type": mechanism_type,
            "mechanism_name": mechanism_name,
            "required_percentage": float(required_percentage),
            "mechanism_description": mechanism_description,
            "updated_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        # 更新或插入设置
        cols["settings"].update_one(
            {"session_id": session_id},
            {"$set": mechanism_doc},
            upsert=True
        )
        
        return jsonify({"code": 200, "message": "投票机制保存成功"})
        
    except Exception as e:
        return jsonify({"code": 500, "message": f"保存投票机制失败: {str(e)}"}), 500

@app.route('/api/get_real_time_voting_data')
def api_get_real_time_voting_data():
    """获取实时投票数据"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取参与国家
        settings = cols["settings"].find_one({"session_id": session_id})
        if not settings or not settings.get("participants"):
            return jsonify({
                "code": 404,
                "message": "未找到参与国家信息"
            })
        
        participant_ids = [str(pid) for pid in settings["participants"]]
        countries = list(cols["countries"].find({"_id": {"$in": [ObjectId(pid) for pid in participant_ids]}}))
        
        # 获取投票数据
        vote_details = cols["file_vote_details"].find_one({"session_id": session_id})
        votes = vote_details.get("votes", {}) if vote_details else {}
        
        # 处理国家数据
        flag_dir = os.path.join(app.static_folder, "flags")
        country_data = []
        vote_data = {}
        
        for country in countries:
            cid = str(country.get('_id'))
            name = country.get('country_name', '未知国家')
            
            # 处理国旗
            if country.get('flag'):
                filename = os.path.basename(country['flag'])
            elif country.get('code'):
                filename = f"{country['code'].lower()}.png"
            else:
                filename = None
            
            if filename and os.path.exists(os.path.join(flag_dir, filename)):
                flag_url = f"/static/flags/{filename}"
            else:
                flag_url = "/static/flags/default.png"
            
            country_data.append({
                "id": cid,
                "name": name,
                "flag_url": flag_url
            })
            
            # 处理投票数据（这里简化处理，实际可能需要更复杂的逻辑）
            if cid in votes:
                vote_data[cid] = votes[cid]
        
        return jsonify({
            "code": 200,
            "message": "获取实时投票数据成功",
            "data": {
                "countries": country_data,
                "votes": vote_data
            }
        })
        
    except Exception as e:
        return jsonify({"code": 500, "message": f"获取实时投票数据失败: {str(e)}"}), 500

@app.route('/api/finalize_voting', methods=['POST'])
def api_finalize_voting():
    """完成投票"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        
        # 获取数据库集合
        cols = get_cols_by_session(session_id)
        
        # 这里可以添加完成投票的逻辑，比如：
        # 1. 计算最终结果
        # 2. 生成投票报告
        # 3. 更新投票状态
        
        # 简单实现：更新会议状态
        cols["settings"].update_one(
            {"session_id": session_id},
            {"$set": {
                "voting_completed": True,
                "voting_completed_at": datetime.now(UTC).isoformat() + "Z"
            }}
        )
        
        return jsonify({"code": 200, "message": "投票已完成"})
        
    except Exception as e:
        return jsonify({"code": 500, "message": f"完成投票失败: {str(e)}"}), 500

# =========================
# API：提交/读取 文本+附件
# =========================
@app.route('/api/submissions', methods=['GET', 'POST'])
def api_submissions():
    """
    GET: /api/submissions?session_id=xxx  -> 返回当前会期所有提交
    POST: 同上 + form-data {country_id, text, file?}
    附件保存到 /app/static/uploads/ 下，返回 file_name，前端用 /static/uploads/<file_name> 打开
    """
    session_id = request.args.get("session_id", "default")
    cols = get_cols_by_session(session_id)

    if request.method == 'GET':
        cur = cols["submissions"].find({"session_id": session_id}).sort("created_at", -1)
        data = []
        for s in cur:
            data.append({
                "country_id": s.get("country_id"),
                "text": s.get("text", ""),
                "file_name": s.get("file_name"),
                "created_at": s.get("created_at"),
            })
        return jsonify({"code": 200, "data": data})

    # POST
    form = request.form
    files = request.files

    country_id = (form.get("country_id") or "").strip()
    text = (form.get("text") or "").strip()
    
    # 检查是否有文件上传
    file_obj = files.get("file")
    has_file = file_obj and file_obj.filename
    
    # 如果既没有文本也没有文件，则返回错误
    if not country_id or (not text and not has_file):
        return jsonify({"code": 400, "msg": "country_id 缺失或需要提供文本内容或文件"}), 400

    # 兜底校验：该会期该国家只能提交一次
    exists = cols["submissions"].find_one({"country_id": country_id, "session_id": session_id})
    if exists:
        return jsonify({"code": 409, "msg": "该国家已提交，不能重复提交"}), 409

    # 处理附件（可选）
    saved_name = None
    if has_file:
        ALLOWED = {".pdf", ".doc", ".docx"}
        ext = Path(file_obj.filename).suffix.lower()
        if ext not in ALLOWED:
            return jsonify({"code": 400, "msg": "不支持的附件类型"}), 400

        upload_dir = Path(current_app.static_folder) / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now(UTC).strftime("%Y%m%d%H%M%S%f")
        saved_name = f"{ts}_{secure_filename(file_obj.filename)}"
        file_obj.save(upload_dir / saved_name)

    record = {
        "country_id": country_id,
        "session_id": session_id,
        "text": text,
        "file_name": saved_name,
        "created_at": datetime.now(UTC).isoformat() + "Z"
    }
    result = cols["submissions"].insert_one(record)
    
    # 返回可序列化的数据
    return_data = {
        "country_id": str(record["country_id"]),
        "session_id": record["session_id"],
        "text": record["text"],
        "file_name": record["file_name"],
        "created_at": record["created_at"]
    }
    return jsonify({"code": 200, "data": return_data})

@app.route('/api/clear_files', methods=['POST'])
def api_clear_files():
    """清理指定会期的所有文件"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取该会期的所有提交记录
        submissions = list(cols["submissions"].find({"session_id": session_id}))
        
        # 删除文件
        upload_dir = Path(current_app.static_folder) / "uploads"
        deleted_files = []
        
        for submission in submissions:
            file_name = submission.get("file_name")
            if file_name:
                file_path = upload_dir / file_name
                if file_path.exists():
                    file_path.unlink()
                    deleted_files.append(file_name)
        
        # 删除数据库记录
        result = cols["submissions"].delete_many({"session_id": session_id})
        
        return jsonify({
            "code": 200, 
            "message": f"清理成功，删除了 {len(deleted_files)} 个文件，{result.deleted_count} 条记录"
        })
        
    except Exception as e:
        print(f"清理文件时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"清理失败: {str(e)}"}), 500

@app.route('/api/save_vote_results', methods=['POST'])
def api_save_vote_results():
    """保存投票结果"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        motion_country_id = data.get("motion_country_id", "")
        motion_country_name = data.get("motion_country_name", "")
        votes = data.get("votes", {})
        motion_text = data.get("motion_text", "")
        
        # 计算投票统计
        agree_count = 0
        disagree_count = 0
        abstain_count = 0
        
        for vote_type in votes.values():
            if vote_type == "agree":
                agree_count += 1
            elif vote_type == "disagree":
                disagree_count += 1
            elif vote_type == "abstain":
                abstain_count += 1
        
        # 判断是否通过（同意票多于不同意票）
        is_passed = agree_count > disagree_count
        
        # 保存到数据库
        vote_record = {
            "session_id": session_id,
            "motion_country_id": motion_country_id,
            "motion_country_name": motion_country_name,
            "motion_text": motion_text,
            "votes": votes,
            "agree_count": agree_count,
            "disagree_count": disagree_count,
            "abstain_count": abstain_count,
            "total_count": len(votes),
            "is_passed": is_passed,
            "created_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        # 保存到投票结果集合
        col_vote_results = db["vote_results"]
        col_vote_results.insert_one(vote_record)
        
        # 更新原始提交记录，标记投票结果
        vote_status = "passed" if is_passed else "failed"
        col_submissions.update_one(
            {"session_id": session_id, "country_id": motion_country_id},
            {"$set": {
                "vote_passed": is_passed, 
                "vote_status": vote_status,
                "vote_updated_at": datetime.now(UTC).isoformat() + "Z",
                "vote_agree_count": agree_count,
                "vote_disagree_count": disagree_count,
                "vote_abstain_count": abstain_count
            }}
        )
        
        return jsonify({
            "code": 200,
            "message": "投票结果保存成功",
            "data": {
                "agree_count": agree_count,
                "disagree_count": disagree_count,
                "abstain_count": abstain_count,
                "is_passed": is_passed
            }
        })
        
    except Exception as e:
        print(f"保存投票结果时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/save_file_vote_results', methods=['POST'])
def api_save_file_vote_results():
    """保存文件投票结果"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        votes = data.get("votes", {})
        files = data.get("files", [])
        
        # 获取文件分配信息
        file_assignments = col_file_assignments.find_one({"session_id": session_id})
        assignments = file_assignments.get("assignments", {}) if file_assignments else {}
        
        # 获取到场国家列表
        arrived_rollcall = list(col_rollcall.find({"session_id": session_id, "arrived": True}))
        arrived_countries = [str(doc.get("country_id")) for doc in arrived_rollcall]
        
        results = []
        
        # 处理每个文件的投票结果
        for file in files:
            file_id = file.get("file_id")
            if file_id not in votes:
                continue
                
            # 获取分配的国家
            assigned_country_id = None
            for country_id, assigned_file_id in assignments.items():
                if assigned_file_id == file_id:
                    assigned_country_id = country_id
                    break
            
            if not assigned_country_id:
                continue
            
            # 计算投票统计
            agree_count = 0
            disagree_count = 0
            abstain_count = 0
            
            # 获取该文件的投票结果
            file_votes = votes.get(file_id, {})
            
            # 如果没有传入投票数据，从数据库获取
            if not file_votes:
                col_file_vote_details = db["file_vote_details"]
                vote_details = list(col_file_vote_details.find({
                    "session_id": session_id,
                    "file_id": file_id
                }))
                file_votes = {detail.get("country_id"): detail.get("vote_result") for detail in vote_details}
            
            for country_id in arrived_countries:
                if country_id == assigned_country_id:
                    continue  # 跳过动议国家本身
                    
                # 获取该国家对该文件的投票
                country_vote = file_votes.get(country_id)
                if country_vote:
                    if country_vote == 'agree':
                        agree_count += 1
                    elif country_vote == 'disagree':
                        disagree_count += 1
                    elif country_vote == 'abstain':
                        abstain_count += 1
            
            # 判断是否通过（同意票多于不同意票）
            is_passed = agree_count > disagree_count
            
            # 获取国家信息
            country_info = col_countries.find_one({"_id": ObjectId(assigned_country_id)})
            country_name = country_info.get("name", "未知国家") if country_info else "未知国家"
            
            # 保存投票结果到数据库
            vote_record = {
                "session_id": session_id,
                "file_id": file_id,
                "file_name": file.get("file_name", ""),
                "country_id": assigned_country_id,
                "country_name": country_name,
                "votes": file_votes,  # 使用实际的投票数据
                "agree_count": agree_count,
                "disagree_count": disagree_count,
                "abstain_count": abstain_count,
                "total_count": len(arrived_countries) - 1,  # 减去动议国家
                "is_passed": is_passed,
                "created_at": datetime.now(UTC).isoformat() + "Z"
            }
            
            # 保存到文件投票结果集合
            col_file_vote_results = db["file_vote_results"]
            col_file_vote_results.insert_one(vote_record)
            
            # 更新临时文件状态
            col_temp_files.update_one(
                {"session_id": session_id, "file_id": file_id},
                {"$set": {
                    "vote_passed": is_passed,
                    "vote_status": "passed" if is_passed else "failed",
                    "vote_updated_at": datetime.now(UTC).isoformat() + "Z",
                    "vote_agree_count": agree_count,
                    "vote_disagree_count": disagree_count,
                    "vote_abstain_count": abstain_count
                }}
            )
            
            results.append({
                "file_id": file_id,
                "file_name": file.get("file_name", ""),
                "country_id": assigned_country_id,
                "country_name": country_name,
                "agree_count": agree_count,
                "disagree_count": disagree_count,
                "abstain_count": abstain_count,
                "is_passed": is_passed
            })
        
        return jsonify({
            "code": 200,
            "message": "文件投票结果保存成功",
            "data": results
        })
        
    except Exception as e:
        print(f"保存文件投票结果时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/get_passed_files', methods=['GET'])
def api_get_passed_files():
    """获取通过投票的文件列表"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 获取通过投票的临时文件
        passed_files = list(col_temp_files.find({
            "session_id": session_id,
            "vote_passed": True
        }))
        
        # 获取文件分配信息
        file_assignments = col_file_assignments.find_one({"session_id": session_id})
        assignments = file_assignments.get("assignments", {}) if file_assignments else {}
        
        # 格式化数据
        data = []
        for file in passed_files:
            # 获取分配的国家
            assigned_country_id = None
            for country_id, assigned_file_id in assignments.items():
                if assigned_file_id == file.get("file_id"):
                    assigned_country_id = country_id
                    break
            
            if assigned_country_id:
                country_info = col_countries.find_one({"_id": ObjectId(assigned_country_id)})
                country_name = country_info.get("name", "未知国家") if country_info else "未知国家"
                
                data.append({
                    "file_id": file.get("file_id"),
                    "file_name": file.get("file_name", ""),
                    "country_id": assigned_country_id,
                    "country_name": country_name,
                    "extracted_text": file.get("extracted_text", ""),
                    "created_at": file.get("created_at"),
                    "vote_updated_at": file.get("vote_updated_at")
                })
        
        return jsonify({
            "code": 200,
            "message": "获取通过投票的文件成功",
            "data": data
        })
        
    except Exception as e:
        print(f"获取通过投票的文件时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

@app.route('/api/get_passed_submissions', methods=['GET'])
def api_get_passed_submissions():
    """获取通过投票的文件列表（优化版：支持passed_files集合）"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        print(f"\n🔍 获取通过投票的文件，session_id: {session_id}")
        
        # 🔥 优先从passed_files集合获取（新版本）
        passed_files = list(cols["db"]["passed_files"].find({
            "session_id": session_id,
            "status": "passed"
        }))
        
        print(f"📁 从passed_files集合找到: {len(passed_files)} 个文件")
        
        # 格式化数据
        data = []
        
        if passed_files:
            # 从passed_files集合读取
            for pf in passed_files:
                country_id = pf.get("country_id", "")
                
                # 尝试获取国家名称
                country_name = ""
                if country_id:
                    try:
                        country = MASTER_COUNTRIES.find_one({"_id": ObjectId(country_id)})
                        if country:
                            country_name = country.get("name", country.get("country_name", ""))
                    except:
                        country_name = country_id
                
                data.append({
                    "country_id": str(country_id),
                    "country_name": country_name,
                    "text": "",  # passed_files不存储text
                    "file_name": pf.get("file_name", ""),
                    "original_name": pf.get("original_name", ""),
                    "created_at": pf.get("passed_at", ""),
                    "vote_agree": pf.get("vote_agree", 0),
                    "vote_disagree": pf.get("vote_disagree", 0),
                    "vote_abstain": pf.get("vote_abstain", 0)
                })
                print(f"  ✅ {country_name}: {pf.get('original_name', '')}")
        else:
            # 回退到submissions集合（旧版本兼容）
            print(f"⚠️  passed_files为空，尝试从submissions获取...")
            passed_submissions = list(cols["submissions"].find({
                "session_id": session_id,
                "vote_passed": True
            }))
            
            print(f"📄 从submissions找到: {len(passed_submissions)} 个文件")
            
            for submission in passed_submissions:
                country_id = submission.get("country_id", "")
                country_name = ""
                
                if country_id:
                    try:
                        country = MASTER_COUNTRIES.find_one({"_id": ObjectId(country_id)})
                        if country:
                            country_name = country.get("name", country.get("country_name", ""))
                    except:
                        country_name = str(country_id)
                
                data.append({
                    "country_id": str(country_id),
                    "country_name": country_name,
                    "text": submission.get("text", ""),
                    "file_name": submission.get("file_name"),
                    "created_at": submission.get("created_at"),
                    "vote_updated_at": submission.get("vote_updated_at")
                })
        
        print(f"✅ 返回 {len(data)} 个通过投票的文件\n")
        
        return jsonify({
            "code": 200,
            "message": "获取通过投票的文件成功",
            "data": data
        })
        
    except Exception as e:
        print(f"❌ 获取通过投票的文件时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

@app.route('/api/generate_consensus_declaration', methods=['POST'])
def api_generate_consensus_declaration():
    """基于通过投票的文件生成共同宣言"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        
        # 优先获取通过投票的临时文件
        passed_files = list(col_temp_files.find({
            "session_id": session_id,
            "vote_passed": True
        }))
        
        # 如果没有临时文件，则获取通过投票的提交
        if not passed_files:
            passed_submissions = list(col_submissions.find({
                "session_id": session_id,
                "vote_passed": True
            }))
            
            if not passed_submissions:
                return jsonify({
                    "code": 400,
                    "message": "没有通过投票的文件，无法生成共同宣言"
                }), 400
            
            # 提取所有文本内容
            texts = [submission.get("text", "") for submission in passed_submissions]
        else:
            # 使用临时文件的文本内容
            texts = [file.get("extracted_text", "") for file in passed_files]
        
        # 使用本地文本相似度分析和生成
        declaration = generate_consensus_declaration_local(texts) #2845行,本地生成
        
        # 保存生成的宣言
        declaration_record = {
            "session_id": session_id,
            "declaration": declaration,
            "source_submissions": [str(s.get("country_id")) for s in passed_submissions],
            "created_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        col_declarations = db["declarations"]
        col_declarations.insert_one(declaration_record)
        
        return jsonify({
            "code": 200,
            "message": "共同宣言生成成功",
            "data": {
                "declaration": declaration,
                "source_count": len(passed_submissions)
            }
        })
        
    except Exception as e:
        print(f"生成共同宣言时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"生成失败: {str(e)}"}), 500

def generate_consensus_declaration_local(texts):
    """使用本地方法生成共同宣言"""
    try:
        # 简单的文本相似度分析
        common_keywords = extract_common_keywords(texts)
        
        # 基于关键词和文本内容生成宣言
        declaration = create_declaration_from_keywords(texts, common_keywords)
        
        return declaration
        
    except Exception as e:
        print(f"本地生成宣言时出错: {str(e)}")
        # 返回一个基本的宣言模板
        return create_basic_declaration(texts)

def extract_common_keywords(texts, top_n=10):
    """提取文本中的共同关键词"""
    try:
        import jieba
        from collections import Counter
        
        # 中文停用词
        stop_words = {
            '的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这',
            '对', '等', '中', '为', '以', '及', '与', '或', '但', '而', '如果', '因为', '所以', '然后', '同时', '另外', '此外', '总之', '因此', '然而', '不过', '虽然', '尽管',
            '关于', '对于', '由于', '通过', '根据', '按照', '基于', '鉴于', '考虑到', '注意到', '强调', '指出', '认为', '表示', '提出', '建议', '支持', '反对', '同意', '不同意'
        }
        
        all_words = []
        for text in texts:
            # 使用jieba进行中文分词
            words = jieba.cut(text)
            words = [word.strip() for word in words if len(word.strip()) > 1 and word.strip() not in stop_words]
            all_words.extend(words)
        
        # 统计词频
        word_counts = Counter(all_words)
        common_keywords = [word for word, count in word_counts.most_common(top_n)]
        
        return common_keywords
        
    except ImportError:
        # 如果没有jieba，使用简单的正则表达式方法
        import re
        from collections import Counter
        
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'}
        
        all_words = []
        for text in texts:
            words = re.findall(r'[\u4e00-\u9fff]+', text)
            words = [word for word in words if len(word) > 1 and word not in stop_words]
            all_words.extend(words)
        
        word_counts = Counter(all_words)
        common_keywords = [word for word, count in word_counts.most_common(top_n)]
        
        return common_keywords

def create_declaration_from_keywords(texts, keywords):
    """基于关键词创建宣言"""
    try:
        # 提取关键句子
        key_sentences = []
        for text in texts:
            sentences = text.split('。')
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) > 10:  # 过滤太短的句子
                    # 计算句子与关键词的相关性
                    relevance_score = sum(1 for keyword in keywords[:5] if keyword in sentence)
                    if relevance_score > 0:
                        key_sentences.append((sentence, relevance_score))
        
        # 按相关性排序并去重
        key_sentences.sort(key=lambda x: x[1], reverse=True)
        unique_sentences = []
        seen_sentences = set()
        for sentence, score in key_sentences:
            # 简单的去重（基于前20个字符）
            sentence_key = sentence[:20]
            if sentence_key not in seen_sentences:
                unique_sentences.append(sentence)
                seen_sentences.add(sentence_key)
            if len(unique_sentences) >= 5:
                break
        
        # 如果没有找到相关句子，使用一些通用句子
        if not unique_sentences:
            unique_sentences = [
                "各方在相关议题上存在广泛共识",
                "支持通过对话和合作解决分歧",
                "致力于推动相关领域的进展"
            ]
        
        # 分析文本主题
        theme_keywords = [kw for kw in keywords if len(kw) > 2][:3]
        theme_text = "、".join(theme_keywords) if theme_keywords else "相关议题"#在一连串的词里面选超过两字节的前三个出现的作为主题
        
        # 构建宣言
        declaration_parts = [
            f"基于各国代表关于{theme_text}的立场文件和投票表决结果，我们达成以下共识：",
            "",
            "一、共同立场",
            *[f"• {sentence}" for sentence in unique_sentences],
            "",
            "二、合作承诺",
            "• 各方承诺在相关领域加强合作与交流",
            "• 支持建立有效的协调机制",
            "• 致力于推动相关议题的进展",
            "",
            "三、后续行动",
            "• 建立定期磋商机制",
            "• 制定具体实施方案",
            "• 定期评估合作进展",
            "",
            "本宣言体现了各方的共同意愿和合作精神，将为相关领域的合作奠定坚实基础。"
        ]
        
        return "\n".join(declaration_parts)
        
    except Exception as e:
        print(f"创建宣言时出错: {str(e)}")
        return create_basic_declaration(texts)

def create_basic_declaration(texts):
    """创建基本宣言模板"""
    return """基于各国代表的立场文件和投票表决结果，我们达成以下共识：

一、共同立场
• 各方在相关议题上存在广泛共识
• 支持通过对话和合作解决分歧
• 致力于推动相关领域的进展

二、合作承诺
• 各方承诺在相关领域加强合作与交流
• 支持建立有效的协调机制
• 致力于推动相关议题的进展

三、后续行动
• 建立定期磋商机制
• 制定具体实施方案
• 定期评估合作进展

本宣言体现了各方的共同意愿和合作精神，将为相关领域的合作奠定坚实基础。"""

@app.route('/api/clear_declarations', methods=['POST'])
def api_clear_declarations():
    """清理指定会期的所有宣言记录"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 删除数据库中的宣言记录
        result = db.declarations.delete_many({"session_id": session_id})
        
        return jsonify({
            "code": 200, 
            "message": f"清理成功，删除了 {result.deleted_count} 条宣言记录"
        })
        
    except Exception as e:
        print(f"清理宣言记录时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"清理失败: {str(e)}"}), 500

@app.route('/api/clear_rollcall', methods=['POST'])
def api_clear_rollcall():
    """清理指定会期的所有点名记录"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 删除数据库中的点名记录
        result = col_rollcall.delete_many({"session_id": session_id})
        
        return jsonify({
            "code": 200, 
            "message": f"清理成功，删除了 {result.deleted_count} 条点名记录"
        })
        
    except Exception as e:
        print(f"清理点名记录时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"清理失败: {str(e)}"}), 500

@app.route('/api/clear_attending_countries', methods=['POST'])
def api_clear_attending_countries():
    """清理指定会期的应出席国家设置"""
    try:
        session_id = request.args.get("session_id", "default")

        # 清空应出席国家列表
        result = col_settings.update_one(
            {"session_id": session_id},
            {"$set": {"participants": []}},
            upsert=True
        )

    except Exception as e:
        print(f"清理点名记录时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"清理失败: {str(e)}"}), 500

# =========================
# 文件上传管理API
# =========================

@app.route('/api/upload_temp_files', methods=['POST'])
def api_upload_temp_files():
    """上传临时文件"""
    try:
        session_id = request.args.get("session_id", "default")
        files = request.files.getlist('files')
        
        if not files:
            return jsonify({"code": 400, "message": "没有上传文件"}), 400
        
        uploaded_files = []
        upload_dir = Path(current_app.static_folder) / "temp_uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        for file in files:
            if file.filename:
                # 生成唯一文件名
                ts = datetime.now(UTC).strftime("%Y%m%d%H%M%S%f")
                filename = f"{ts}_{secure_filename(file.filename)}"
                file_path = upload_dir / filename
                
                # 保存文件
                file.save(file_path)
                
                # 保存到数据库
                file_record = {
                    "session_id": session_id,
                    "original_name": file.filename,
                    "saved_name": filename,
                    "file_path": str(file_path),
                    "file_size": file.content_length or 0,
                    "content_type": file.content_type,
                    "uploaded_at": datetime.now(UTC).isoformat() + "Z"
                }
                
                result = col_temp_files.insert_one(file_record)
                file_record["_id"] = str(result.inserted_id)
                uploaded_files.append(file_record)
        
        return jsonify({
            "code": 200,
            "message": f"成功上传 {len(uploaded_files)} 个文件",
            "data": uploaded_files
        })
        
    except Exception as e:
        print(f"上传临时文件时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"上传失败: {str(e)}"}), 500

@app.route('/api/save_file_assignments', methods=['POST'])
def api_save_file_assignments():
    """保存文件分配"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        files = data.get("files", [])
        assignments = data.get("assignments", {})
        
        # 保存文件信息
        for file_info in files:
            file_record = {
                "session_id": session_id,
                "file_id": file_info.get("id"),
                "file_name": file_info.get("name"),
                "file_size": file_info.get("size"),
                "file_type": file_info.get("type"),
                "uploaded_at": datetime.now(UTC).isoformat() + "Z"
            }
            
            col_temp_files.update_one(
                {"session_id": session_id, "file_id": file_info.get("id")},
                {"$set": file_record},
                upsert=True
            )
        
        # 保存分配信息
        assignment_record = {
            "session_id": session_id,
            "assignments": assignments,
            "created_at": datetime.now(UTC).isoformat() + "Z"
        }
        
        col_file_assignments.update_one(
            {"session_id": session_id},
            {"$set": assignment_record},
            upsert=True
        )
        
        # 同时保存到投票文件表
        vote_files = []
        for country_id, file_id in assignments.items():
            # 获取文件信息
            file_info = next((f for f in files if f.get("id") == file_id), None)
            if file_info:
                # 获取国家信息
                country_info = col_countries.find_one({"_id": ObjectId(country_id)})
                country_name = country_info.get("country_name", "未知国家") if country_info else "未知国家"
                
                # 处理国旗文件名
                if country_info and country_info.get('flag'):
                    country_flag = os.path.basename(country_info['flag'])
                elif country_info and country_info.get('code'):
                    country_flag = f"{country_info['code'].lower()}.png"
                else:
                    country_flag = "default.png"
                
                # 获取文件的提取文本
                temp_file = col_temp_files.find_one({"session_id": session_id, "file_id": file_id})
                extracted_text = temp_file.get("extracted_text", "") if temp_file else ""
                
                vote_file_record = {
                    "session_id": session_id,
                    "file_id": file_id,
                    "file_name": file_info.get("name", ""),
                    "file_path": file_info.get("path", ""),
                    "extracted_text": extracted_text,
                    "country_id": country_id,
                    "country_name": country_name,
                    "country_flag": country_flag,
                    "vote_status": "pending",  # pending, voted, passed, failed
                    "vote_result": None,  # agree, disagree, abstain
                    "created_at": datetime.now(UTC).isoformat() + "Z"
                }
                vote_files.append(vote_file_record)
        
        # 清空旧的投票文件记录，插入新的
        col_vote_files.delete_many({"session_id": session_id})
        if vote_files:
            col_vote_files.insert_many(vote_files)
        
        return jsonify({
            "code": 200,
            "message": "文件分配保存成功",
            "data": {
                "files_count": len(files),
                "assignments_count": len(assignments)
            }
        })
        
    except Exception as e:
        print(f"保存文件分配时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/get_temp_files', methods=['GET'])
def api_get_temp_files():
    """获取临时文件列表"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 获取文件列表
        files = list(col_temp_files.find({"session_id": session_id}))
        
        # 获取分配信息
        assignment_doc = col_file_assignments.find_one({"session_id": session_id})
        assignments = assignment_doc.get("assignments", {}) if assignment_doc else {}
        
        # 格式化数据
        formatted_files = []
        for file_doc in files:
            formatted_files.append({
                "id": file_doc.get("file_id", str(file_doc["_id"])),
                "name": file_doc.get("file_name", file_doc.get("original_name", "")),
                "size": file_doc.get("file_size", 0),
                "type": file_doc.get("content_type", ""),
                "uploaded_at": file_doc.get("uploaded_at")
            })
        
        return jsonify({
            "code": 200,
            "data": {
                "files": formatted_files,
                "assignments": assignments
            }
        })
        
    except Exception as e:
        print(f"获取临时文件时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

@app.route('/api/clear_temp_files', methods=['POST'])
def api_clear_temp_files():
    """清理临时文件"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 获取要删除的文件
        files = list(col_temp_files.find({"session_id": session_id}))
        
        # 删除物理文件
        upload_dir = Path(current_app.static_folder) / "temp_uploads"
        deleted_files = []
        
        for file_doc in files:
            file_path = upload_dir / file_doc.get("saved_name", "")
            if file_path.exists():
                file_path.unlink()
                deleted_files.append(file_doc.get("saved_name"))
        
        # 删除数据库记录
        col_temp_files.delete_many({"session_id": session_id})
        col_file_assignments.delete_many({"session_id": session_id})
        
        return jsonify({
            "code": 200,
            "message": f"清理成功，删除了 {len(deleted_files)} 个文件"
        })
        
    except Exception as e:
        print(f"清理临时文件时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"清理失败: {str(e)}"}), 500

@app.route('/api/get_vote_files', methods=['GET'])
def api_get_vote_files():
    """获取投票文件列表"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 从投票文件表获取数据
        vote_files = list(col_vote_files.find({"session_id": session_id}))
        
        # 格式化数据
        data = []
        for file in vote_files:
            data.append({
                "file_id": file.get("file_id"),
                "file_name": file.get("file_name", ""),
                "extracted_text": file.get("extracted_text", ""),
                "country_id": file.get("country_id"),
                "country_name": file.get("country_name", ""),
                "country_flag": file.get("country_flag", "default.png"),
                "vote_status": file.get("vote_status", "pending"),
                "vote_result": file.get("vote_result"),
                "created_at": file.get("created_at")
            })
        
        return jsonify({
            "code": 200,
            "message": "获取投票文件列表成功",
            "data": data
        })
        
    except Exception as e:
        print(f"获取投票文件列表时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

@app.route('/api/save_single_vote', methods=['POST'])
def api_save_single_vote():
    """保存单个文件的投票结果"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        file_id = data.get("file_id")
        vote_result = data.get("vote_result")  # agree, disagree, abstain
        
        if not file_id or not vote_result:
            return jsonify({"code": 400, "message": "缺少必要参数"}), 400
        
        # 更新投票文件记录
        result = col_vote_files.update_one(
            {"session_id": session_id, "file_id": file_id},
            {"$set": {
                "vote_status": "voted",
                "vote_result": vote_result,
                "voted_at": datetime.now(UTC).isoformat() + "Z"
            }}
        )
        
        if result.modified_count > 0:
            return jsonify({
                "code": 200,
                "message": "投票结果保存成功"
            })
        else:
            return jsonify({
                "code": 404,
                "message": "未找到对应的文件"
            }), 404
        
    except Exception as e:
        print(f"保存单个投票结果时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/save_batch_votes', methods=['POST'])
def api_save_batch_votes():
    """保存批量投票结果 - 每个国家对每个文件的投票"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        votes_data = data.get("votes", {})  # {file_id: {country_id: vote_result}}
        
        if not votes_data:
            return jsonify({"code": 400, "message": "投票数据不能为空"}), 400
        
        # 创建或更新投票结果集合
        col_file_vote_details = db["file_vote_details"]
        
        saved_count = 0
        for file_id, country_votes in votes_data.items():
            for country_id, vote_result in country_votes.items():
                # 保存每个国家对每个文件的投票详情
                vote_detail = {
                    "session_id": session_id,
                    "file_id": file_id,
                    "country_id": country_id,
                    "vote_result": vote_result,
                    "voted_at": datetime.now(UTC).isoformat() + "Z"
                }
                
                # 使用upsert避免重复
                col_file_vote_details.update_one(
                    {
                        "session_id": session_id,
                        "file_id": file_id,
                        "country_id": country_id
                    },
                    {"$set": vote_detail},
                    upsert=True
                )
                saved_count += 1
        
        return jsonify({
            "code": 200,
            "message": f"成功保存 {saved_count} 个投票记录"
        })
        
    except Exception as e:
        print(f"保存批量投票结果时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"保存失败: {str(e)}"}), 500

@app.route('/api/get_file_vote_details_by_session', methods=['GET'])
def api_get_file_vote_details_by_session():
    """获取文件投票详情（使用会话特定的数据库）"""
    try:
        session_id = request.args.get("session_id", "default")

        # 使用与投票提交相同的数据库连接方式
        cols = get_cols_by_session(session_id)

        # 获取投票详情
        vote_details = list(cols["db"]["file_vote_details"].find({"session_id": session_id}))

        # 转换为主席投票监控页面期望的格式：数组格式
        formatted_votes = []
        for detail in vote_details:
            formatted_votes.append({
                "country_id": detail.get("country_id"),
                "file_id": detail.get("file_id"),
                "vote_result": detail.get("vote_result")
            })

        return jsonify({
            "code": 200,
            "message": "获取投票详情成功",
            "data": formatted_votes
        })

    except Exception as e:
        print(f"获取文件投票详情时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

@app.route('/api/get_file_vote_details', methods=['GET'])
def api_get_file_vote_details():
    """获取文件投票详情（主数据库版本，向后兼容）"""
    try:
        session_id = request.args.get("session_id", "default")

        # 获取投票详情
        col_file_vote_details = db["file_vote_details"]
        vote_details = list(col_file_vote_details.find({"session_id": session_id}))

        # 按文件ID分组
        votes_by_file = {}
        for detail in vote_details:
            file_id = detail.get("file_id")
            if file_id not in votes_by_file:
                votes_by_file[file_id] = {}
            votes_by_file[file_id][detail.get("country_id")] = detail.get("vote_result")

        return jsonify({
            "code": 200,
            "message": "获取投票详情成功",
            "data": votes_by_file
        })

    except Exception as e:
        print(f"获取文件投票详情时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

@app.route('/api/debug_vote_files', methods=['GET'])
def api_debug_vote_files():
    """调试：查看投票文件表的数据"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 查看所有相关表的数据
        temp_files = list(col_temp_files.find({"session_id": session_id}))
        file_assignments = list(col_file_assignments.find({"session_id": session_id}))
        vote_files = list(col_vote_files.find({"session_id": session_id}))
        
        return jsonify({
            "code": 200,
            "message": "调试数据获取成功",
            "data": {
                "temp_files_count": len(temp_files),
                "temp_files": temp_files,
                "file_assignments_count": len(file_assignments),
                "file_assignments": file_assignments,
                "vote_files_count": len(vote_files),
                "vote_files": vote_files
            }
        })
        
    except Exception as e:
        print(f"调试数据获取时出错: {str(e)}")
        return jsonify({"code": 500, "message": f"获取失败: {str(e)}"}), 500

# =========================
# 你之前的 flag 文件直出接口（看起来已不用）
# =========================
@app.route('/flags/<filename>')
def get_flag(filename):
    # 这条看起来没用了（路径也写成了 frotend），保留不影响
    return send_from_directory("frotend/public/flags", filename)

# =========================
# 共同宣言功能
# =========================
@app.route('/declaration')
@app.route('/country-declaration')
def declaration_page():
    """共同宣言页面 / 与会国共同宣言页面"""
    session_id = request.args.get("session_id", "default")
    sdoc = col_settings.find_one({"session_id": session_id}) or {}
    committee = sdoc.get("committee_name", " ")
    agenda = sdoc.get("agenda", " ")

    # 获取已提交文件的国家
    submitted_files = list(col_submissions.find({"session_id": session_id}))

    # 统计提交情况
    submitted_countries = []
    total_countries = []

    # 获取所有参与国家
    if sdoc.get("participants"):
        for country_id in sdoc["participants"]:
            country = col_countries.find_one({"_id": country_id})
            if country:
                total_countries.append(country.get("country_name", "未知国家"))

    # 获取已提交的国家
    for submission in submitted_files:
        country_id = submission.get("country_id")
        if country_id:
            country = col_countries.find_one({"_id": country_id})
            if country:
                submitted_countries.append(country.get("country_name", "未知国家"))

    return render_template(
        'declaration.html',
        committee_name=committee,
        agenda=agenda,
        session_id=session_id,
        submitted_files=submitted_files,
        submitted_countries=submitted_countries,
        total_countries=total_countries,
        topic=agenda
    )

@app.route('/api/check_submissions')
def check_submissions():
    """检查是否有提交的文件"""
    try:
        session_id = request.args.get("session_id", "default")
        submitted_files = list(col_submissions.find({"session_id": session_id}))
        
        return jsonify({
            "has_submissions": len(submitted_files) > 0,
            "count": len(submitted_files)
        })
    except Exception as e:
        print(f"检查提交文件时出错: {str(e)}")
        return jsonify({"error": f"检查失败: {str(e)}"}), 500

@app.route('/api/generate_declaration', methods=['POST'])
def generate_declaration():
    """生成共同宣言的API - 基于投票通过的文件和文本内容"""
    try:
        session_id = request.args.get("session_id", "default")
        # 尝试从请求体获取参数
        try:
            request_data = request.get_json() or {}
        except:
            request_data = {}

        print(f"\n{'='*60}")
        print(f"🚀 生成宣言API被调用，session_id: {session_id}")
        print(f"📋 请求参数: {request_data}")

        cols = get_cols_by_session(session_id)

        # 初始化变量
        generation_method = "未知"

        # 【优化】优先从passed_files集合获取通过投票的文件
        passed_files = list(cols["db"]["passed_files"].find({
            "session_id": session_id,
            "status": "passed"
        }))
        print(f"📁 从passed_files集合找到通过的文件数量: {len(passed_files)}")
        
        # 如果passed_files为空，回退到submissions集合
        if not passed_files:
            print("⚠️  passed_files为空，尝试从submissions集合获取...")
            submitted_files = list(cols["submissions"].find({
                "session_id": session_id,
                "vote_passed": True
            }))
            print(f"📄 从submissions集合找到投票通过的文件数量: {len(submitted_files)}")
        else:
            submitted_files = []

        if not passed_files and not submitted_files:
            print("❌ 没有找到投票通过的文件")
            return jsonify({"error": "没有找到投票通过的文件，无法生成共同宣言"}), 400
        
        # 准备提交给大模型的数据(提取关键词)
        sdoc = cols["settings"].find_one({"session_id": session_id}) or {}
        topic = sdoc.get("agenda", "未知议题")
        countries_data = []
        
        # 【优化】处理passed_files集合中的文件
        if passed_files:
            print(f"\n📝 开始处理passed_files中的 {len(passed_files)} 个文件...")
            for passed_file in passed_files:
                country_id = passed_file.get("country_id", "")
                file_name = passed_file.get("file_name", "")
                original_name = passed_file.get("original_name", file_name)
                
                # 获取国家信息
                country_name = "未知国家"
                if country_id:
                    # 尝试从ObjectId获取
                    try:
                        country = MASTER_COUNTRIES.find_one({"_id": ObjectId(country_id)})
                        if country:
                            country_name = country.get("name", country.get("country_name", "未知国家"))
                    except:
                        # 如果不是ObjectId，直接使用country_id作为名称
                        country_name = country_id
                
                print(f"\n🌍 处理国家: {country_name}")
                print(f"📄 文件名: {file_name}")
                print(f"📋 原始文件名: {original_name}")
                
                # 从submissions获取手动输入的文本（如果有）
                manual_text = ""
                submission = cols["submissions"].find_one({
                    "session_id": session_id,
                    "country_id": country_id
                })
                if submission:
                    manual_text = submission.get("text", "")
                    print(f"✍️  手动文本长度: {len(manual_text)}")
                
                # 【关键】提取PDF/文档文件的文本内容
                file_text = ""
                if file_name:
                    upload_dir = Path(current_app.static_folder) / "uploads"
                    file_path = upload_dir / file_name
                    print(f"📂 文件路径: {file_path}")
                    print(f"✅ 文件存在: {file_path.exists()}")
                    
                    if file_path.exists():
                        try:
                            file_text = extract_text_from_file(file_path)
                            print(f"📖 提取的原始文件文本长度: {len(file_text) if file_text else 0}")
                            if file_text:
                                file_text = clean_text(file_text)
                                print(f"🧹 清理后的文件文本长度: {len(file_text)}")
                            else:
                                print("⚠️  文件文本提取为空")
                        except Exception as e:
                            print(f"❌ 提取文件文本失败: {str(e)}")
                    else:
                        print(f"⚠️  文件不存在: {file_path}")
                
                # 合并文本内容
                if file_text:
                    combined_text = file_text
                    if manual_text:
                        combined_text += f"\n\n【补充说明】\n{manual_text}"
                else:
                    combined_text = manual_text
                
                print(f"📊 合并后文本长度: {len(combined_text) if combined_text else 0}")
                
                if combined_text:
                    # 不限制长度，保留完整内容给大模型
                    countries_data.append({
                        "country": country_name,
                        "content": combined_text,  # 保留完整内容
                        "file_content": file_text if file_text else "",
                        "manual_content": manual_text if manual_text else "",
                        "file_name": original_name
                    })
                    print(f"✅ 已添加到countries_data: {country_name}")
                else:
                    print(f"⚠️  跳过国家 {country_name}: 没有有效文本内容")
        
        # 【兼容】处理submissions集合中的文件（向后兼容）
        elif submitted_files:
            print(f"\n📝 开始处理submissions中的 {len(submitted_files)} 个文件...")
            for submission in submitted_files:
                country_id = submission.get("country_id")
                if country_id:
                    country = MASTER_COUNTRIES.find_one({"_id": country_id})
                    if country:
                        country_name = country.get("name", country.get("country_name", "未知国家"))
                        print(f"\n🌍 处理国家: {country_name}")
                        
                        manual_text = submission.get("text", "")
                        print(f"✍️  手动文本长度: {len(manual_text)}")
                        
                        file_text = ""
                        file_name = submission.get("file_name")
                        if file_name:
                            upload_dir = Path(current_app.static_folder) / "uploads"
                            file_path = upload_dir / file_name
                            print(f"📂 文件路径: {file_path}")
                            if file_path.exists():
                                file_text = extract_text_from_file(file_path)
                                print(f"📖 提取的文件文本长度: {len(file_text) if file_text else 0}")
                                if file_text:
                                    file_text = clean_text(file_text)
                                    print(f"🧹 清理后文本长度: {len(file_text)}")
                        
                        if file_text:
                            combined_text = file_text
                            if manual_text:
                                combined_text += f"\n\n【补充说明】\n{manual_text}"
                        else:
                            combined_text = manual_text
                        
                        if combined_text:
                            countries_data.append({
                                "country": country_name,
                                "content": combined_text,
                                "file_content": file_text if file_text else "",
                                "manual_content": manual_text if manual_text else ""
                            })
                            print(f"✅ 已添加到countries_data: {country_name}")
                        else:
                            print(f"⚠️  跳过国家 {country_name}: 没有有效文本内容")
        
        if not countries_data:
            print("❌ 没有找到有效的文本内容")
            return jsonify({
                "error": "没有找到有效的文本内容",
                "generation_method": "无数据"
            }), 400
        
        # 调用大模型生成共同宣言
        print(f"\n{'='*60}")
        print(f"🤖 准备调用大模型生成共同宣言")
        print(f"📌 主题: {topic}")
        print(f"🌍 参与国家数量: {len(countries_data)}")
        for i, data in enumerate(countries_data, 1):
            print(f"   {i}. {data['country']} - 文本长度: {len(data.get('content', ''))}")

        print(f"\n🚀 开始调用大模型API...")
        declaration_result = call_llm_for_declaration(topic, countries_data)
        declaration_text = declaration_result.get("text", "")
        generation_method = declaration_result.get("method", "未知")
        print(f"✅ 大模型返回宣言长度: {len(declaration_text) if declaration_text else 0}")
        print(f"🎯 生成方式: {generation_method}")
        print(f"{'='*60}\n")
        
        # 保存生成的宣言到数据库
        declaration_record = {
            "session_id": session_id,
            "topic": topic,
            "declaration": declaration_text,
            "participating_countries": [data["country"] for data in countries_data],
            "generated_at": datetime.now(UTC),
            "status": "draft",
            "analysis_info": {
                "total_countries": len(countries_data),
                "files_processed": len([d for d in countries_data if d.get("file_content")]),
                "manual_text_count": len([d for d in countries_data if d.get("manual_content")])
            }
        }
        
        # 🔥 修复：使用正确的数据库引用
        try:
            cols["db"]["declarations"].insert_one(declaration_record)
            print(f"💾 宣言已保存到数据库")
        except Exception as save_error:
            print(f"⚠️  保存到数据库失败: {save_error}")
            # 即使保存失败，也返回生成的宣言
        
        print(f"\n✅ 共同宣言生成成功！")
        print(f"   - 宣言长度: {len(declaration_text)} 字")
        print(f"   - 参与国家: {len(countries_data)} 个")
        print(f"   - 生成方式: {generation_method}")
        print(f"{'='*60}\n")

        return jsonify({
            "success": True,
            "declaration": declaration_text,
            "participating_countries": [data["country"] for data in countries_data],
            "generation_method": generation_method,
            "analysis_info": declaration_record["analysis_info"]
        })
        
    except Exception as e:
        print(f"❌ 生成共同宣言时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": f"生成宣言失败: {str(e)}",
            "generation_method": "错误回退"
        }), 500

@app.route('/api/save_declaration', methods=['POST'])
def save_declaration():
    """保存最终版本的共同宣言"""
    try:
        session_id = request.args.get("session_id", "default")
        data = request.get_json()
        declaration_text = data.get("declaration", "").strip()
        
        if not declaration_text:
            return jsonify({"error": "宣言内容不能为空"}), 400
        
        # 更新数据库中的宣言状态
        db.declarations.update_one(
            {"session_id": session_id, "status": "draft"},
            {
                "$set": {
                    "declaration": declaration_text,
                    "status": "final",
                    "finalized_at": datetime.now(UTC)
                }
            }
        )
        
        return jsonify({"success": True, "message": "宣言已保存"})
        
    except Exception as e:
        print(f"保存宣言时出错: {str(e)}")
        return jsonify({"error": f"保存失败: {str(e)}"}), 500

@app.route('/api/get_current_motion', methods=['GET'])
def api_get_current_motion():
    """获取当前动议信息"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)

        # 获取当前活跃的动议
        current_motion = cols["db"]["motion_records"].find_one(
            {"session_id": session_id, "status": "active"},
            sort=[("created_at", -1)]  # 获取最新的活跃动议
        )

        if current_motion:
            # 获取动议国家的详细信息
            country_info = cols["countries"].find_one({"id": current_motion.get("country_id", "")})

            return jsonify({
                "code": 200,
                "message": "获取当前动议成功",
                "data": {
                    "country_name": country_info.get("name", "未知国家") if country_info else "未知国家",
                    "country_id": current_motion.get("country_id", ""),
                    "motion_text": current_motion.get("motion_text", ""),
                    "created_at": current_motion.get("created_at"),
                    "status": current_motion.get("status", "active")
                }
            })
        else:
            return jsonify({
                "code": 200,
                "message": "暂无活跃动议",
                "data": None
            })

    except Exception as e:
        return jsonify({"code": 500, "message": "Failed to get current motion: {}".format(str(e))}), 500

@app.route('/api/update_meeting_status', methods=['POST'])
def api_update_meeting_status():
    """更新会议状态"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        status = data.get("status", "active")  # active, paused, ended
        updated_by = data.get("updated_by", "system")

        cols = get_cols_by_session(session_id)

        # 更新或插入会议状态记录
        status_record = {
            "session_id": session_id,
            "status": status,
            "updated_at": datetime.now(),
            "updated_by": updated_by
        }

        # 使用upsert操作
        cols["db"]["meeting_status"].replace_one(
            {"session_id": session_id},
            status_record,
            upsert=True
        )

        return jsonify({
            "code": 200,
            "message": "Meeting status updated to: {}".format(status)
        })

    except Exception as e:
        return jsonify({"code": 500, "message": "Failed to update meeting status: {}".format(str(e))}), 500

@app.route('/api/get_declaration', methods=['GET'])
def get_declaration():
    """获取当前会期的共同宣言"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 获取最新的宣言
        declaration = db.declarations.find_one(
            {"session_id": session_id},
            sort=[("generated_at", -1)]
        )
        
        if declaration:
            return jsonify({
                "success": True,
                "declaration": declaration.get("declaration", ""),
                "status": declaration.get("status", "draft"),
                "participating_countries": declaration.get("participating_countries", []),
                "generated_at": declaration.get("generated_at")
            })
        else:
            return jsonify({"success": False, "message": "暂无宣言"})
            
    except Exception as e:
        print(f"获取宣言时出错: {str(e)}")
        return jsonify({"error": f"获取失败: {str(e)}"}), 500

@app.route('/api/generate_declaration_from_files', methods=['POST'])
def api_generate_declaration_from_files():
    """基于上传的文件生成共同宣言"""
    try:
        # 获取生成方法
        method = request.form.get('method', 'ai')

        # 获取上传的文件
        files = request.files
        if not files:
            return jsonify({"success": False, "error": "没有上传文件"}), 400

        uploaded_files = []
        extracted_texts = []

        # 处理每个上传的文件
        for key in files:
            if key.startswith('file_'):
                file = files[key]
                if file.filename:
                    # 保存文件到临时目录
                    upload_dir = Path(current_app.static_folder) / "temp_uploads"
                    upload_dir.mkdir(parents=True, exist_ok=True)

                    # 生成唯一文件名
                    ts = datetime.now(UTC).strftime("%Y%m%d%H%M%S%f")
                    filename = f"{ts}_{secure_filename(file.filename)}"
                    file_path = upload_dir / filename

                    # 保存文件
                    file.save(file_path)
                    uploaded_files.append({
                        "original_name": file.filename,
                        "saved_name": filename,
                        "file_path": str(file_path),
                        "file_size": file.content_length or 0
                    })

                    # 提取文本内容
                    try:
                        if file.filename.lower().endswith('.pdf'):
                            text = extract_text_from_pdf(file_path)
                        elif file.filename.lower().endswith(('.doc', '.docx')):
                            text = extract_text_from_docx(file_path)
                        else:
                            text = ""

                        if text.strip():
                            extracted_texts.append({
                                "filename": file.filename,
                                "content": text[:3000]  # 限制文本长度
                            })
                    except Exception as e:
                        print(f"提取 {file.filename} 文本失败: {e}")

        if not extracted_texts:
            return jsonify({"success": False, "error": "未能从上传的文件中提取到有效文本内容"}), 400

        # 构建国家数据（假设每个文件代表一个国家）
        countries_data = []
        for i, text_info in enumerate(extracted_texts):
            countries_data.append({
                "country": f"国家{i+1}",  # 暂时用序号代替国家名
                "content": text_info["content"]
            })

        # 生成宣言
        topic = "国际贸易谈判"  # 默认主题

        if method == 'ai':
            try:
                # 尝试使用千问API
                declaration = call_qianwen_api(topic, countries_data)
                generation_method = "通义千问"
            except Exception as e:
                print(f"千问API失败，回退到本地生成: {e}")
                # 回退到本地生成
                texts = [data['content'] for data in countries_data]
                declaration = generate_consensus_declaration_local(texts)
                generation_method = "本地生成（回退）"
        else:
            # 直接使用本地生成
            texts = [data['content'] for data in countries_data]
            declaration = generate_consensus_declaration_local(texts)
            generation_method = "本地生成"

        # 清理临时文件
        for file_info in uploaded_files:
            try:
                file_path = Path(file_info["file_path"])
                if file_path.exists():
                    file_path.unlink()
            except:
                pass

        return jsonify({
            "success": True,
            "method": method,
            "file_count": len(extracted_texts),
            "generation_method": generation_method,
            "declaration": declaration,
            "files_processed": [f["filename"] for f in extracted_texts]
        })

    except Exception as e:
        print(f"生成宣言失败: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

def extract_text_from_pdf(file_path):
    """从PDF文件中提取文本"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        print(f"PDF文本提取失败: {e}")
        return ""

def extract_text_from_docx(file_path):
    """从Word文档中提取文本"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Word文本提取失败: {e}")
        return ""

def call_llm_for_declaration(topic, countries_data):
    """
    调用外部AI API生成共同宣言；失败则回退到本地宣言生成功能。

    返回格式:
    {
        "text": "生成的宣言文本",
        "method": "生成方式"
    }
    """
    try:
        # 调用通义千问API生成宣言
        declaration_text = call_qianwen_api(topic, countries_data)
        return {
            "text": declaration_text,
            "method": "通义千问"
        }

    except Exception as e:
        print(f"千问API调用失败，将回退本地算法: {str(e)}")
        try:
            declaration_text = generate_similarity_based_declaration(topic, countries_data)
            return {
                "text": declaration_text,
                "method": "本地算法（回退）"
            }
        except Exception:
            declaration_text = generate_fallback_declaration(topic, countries_data)
            return {
                "text": declaration_text,
                "method": "模板生成（最终回退）"
            }

def call_qianwen_api(topic, countries_data):
    """
    调用通义千问API生成共同宣言
    使用HTTP请求进行文本生成
    """
    # 通义千问API配置
    # API Key已配置，获取地址：https://dashscope.aliyuncs.com/
    API_KEY = os.getenv('LLM_API_KEY', 'your-llm-api-key-here')
    API_URL = os.getenv('LLM_API_URL', 'https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation')

    # 构建请求数据（千问API格式）
    prompt = build_declaration_prompt(topic, countries_data)

    request_data = {
        "model": os.getenv('LLM_MODEL', 'qwen-turbo'),
        "input": {
            "prompt": prompt
        },
        "parameters": {
            "temperature": 0.3,
            "max_tokens": 2000,
            "top_p": 0.8
        }
    }

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    try:
        # 发送HTTP请求
        response = requests.post(API_URL, json=request_data, headers=headers, timeout=60)

        if response.status_code == 200:
            result = response.json()

            # 尝试多种可能的响应格式
            if result.get("output") and result["output"].get("text"):
                return result["output"]["text"].strip()
            elif result.get("output") and result["output"].get("content"):
                return result["output"]["content"].strip()
            elif result.get("text"):
                return result["text"].strip()
            elif result.get("content"):
                return result["content"].strip()
            else:
                print(f"千问API响应结构: {list(result.keys())}")
                raise Exception("千问API返回格式错误")
        else:
            error_msg = f"千问API请求失败: {response.status_code} - {response.text}"
            print(error_msg)
            print(f"完整响应: {response.text}")
            raise Exception(error_msg)

    except requests.exceptions.RequestException as e:
        print(f"千问API网络请求失败: {str(e)}")
        raise Exception(f"网络请求失败: {str(e)}")
    except Exception as e:
        print(f"千问API调用失败: {str(e)}")
        raise Exception(f"API调用失败: {str(e)}")

def build_declaration_prompt(topic, countries_data):
    """构建千问API的提示词"""
    prompt = f"""你是一名WTO谈判专家与文本分析专家。请基于以下各国提交的文档，生成一份体现最大相似度与共识的共同宣言。

【谈判主题】{topic}

【各国提交内容】
"""
    for i, country_data in enumerate(countries_data, 1):
        prompt += f"\n{i}. {country_data['country']}：\n{country_data['content'][:2000]}"  # 限制长度避免超限

    prompt += """

【生成要求】
1) 先进行相似度分析（共同主题、关键词、代表性表述）。
2) 提取相似度高的关键语句，去重并整合。
3) 生成正式、专业、结构清晰（前言/正文/结论）的共同宣言。
4) 语言中文，800-1200字，尽量保留各国原始表述。
5) 确保宣言体现WTO谈判的专业性和权威性。

【输出】
仅输出共同宣言正文，不要包含任何其他内容。
"""

    return prompt

# 注意：测试函数已移至独立的 test_declaration_api.py 文件中

def generate_similarity_based_declaration(topic, countries_data):
    """
    基于相似度分析生成共同宣言
    直接调用的本地宣言生成功能
    """
    if not countries_data:
        return generate_fallback_declaration(topic, countries_data)

    # 提取所有文本内容，用于本地宣言生成
    texts = [data.get('content', '') for data in countries_data]

    try:
        # 直接调用真正的本地宣言生成函数
        return generate_consensus_declaration_local(texts)
    except Exception as e:
        print(f"本地宣言生成失败，回退到基础模板: {str(e)}")
        return generate_fallback_declaration(topic, countries_data)

# 删除重复的模拟关键词提取函数，使用真正的关键词提取函数（第2861行）

# 删除不再使用的模拟函数，这些功能已被真正的本地宣言生成函数替代

def generate_fallback_declaration(topic, countries_data):
    """生成备用的基础宣言"""
    countries = [data["country"] for data in countries_data]
    countries_str = "、".join(countries)
    
    return f"""
关于{topic}的共同宣言

我们，{countries_str}的代表，在WTO框架下就{topic}议题进行了讨论，达成以下共识：

各方同意在{topic}领域加强合作，推动相关规则的完善和发展。我们将继续通过对话和磋商解决分歧，寻求互利共赢的解决方案。

本宣言体现了各方在此议题上的共同意愿和努力方向。
    """.strip()

# =========================
# 文本提取功能
# =========================
def extract_text_from_pdf(file_path):
    """从PDF文件中提取文本"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"PDF文本提取失败: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    """从Word文档中提取文本"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Word文档文本提取失败: {str(e)}")
        return ""

def extract_text_from_file(file_path):
    """根据文件扩展名提取文本"""
    file_path = Path(file_path)
    if not file_path.exists():
        return ""
    
    file_name = file_path.name.lower()
    ext = file_path.suffix.lower()
    
    # 检查文件扩展名
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.doc', '.docx']:
        return extract_text_from_docx(file_path)
    # 检查文件名中是否包含文件类型标识
    elif '_pdf' in file_name:
        return extract_text_from_pdf(file_path)
    elif '_doc' in file_name or '_docx' in file_name:
        return extract_text_from_docx(file_path)
    else:
        print(f"无法识别的文件类型: {file_name}")
        return ""

def clean_text(text):
    """清理和预处理文本"""
    if not text:
        return ""
    
    # 移除多余的空白字符
    text = re.sub(r'\s+', ' ', text)
    # 移除特殊字符但保留中文、英文、数字和基本标点
    text = re.sub(r'[^\u4e00-\u9fff\w\s.,;:!?()（）【】""''、。，；：！？]', '', text)
    return text.strip()

# =========================
# 启动
# =========================

@app.route('/api/export_declaration_pdf', methods=['POST'])
def export_declaration_pdf():
    """导出共同宣言为PDF文件"""
    try:
        session_id = request.args.get("session_id", "default")
        data = request.get_json()
        declaration_text = data.get("declaration", "").strip()
        
        if not declaration_text:
            return jsonify({"error": "宣言内容不能为空"}), 400
        
        # 获取会议信息
        sdoc = col_settings.find_one({"session_id": session_id}) or {}
        committee_name = sdoc.get("committee_name", "WTO模拟谈判")
        agenda = sdoc.get("agenda", "未指定议题")
        
        # 获取参与国家
        submitted_files = list(col_submissions.find({"session_id": session_id}))
        participating_countries = []
        for submission in submitted_files:
            country_id = submission.get("country_id")
            if country_id:
                country = col_countries.find_one({"_id": country_id})
                if country:
                    participating_countries.append(country.get("country_name", "未知国家"))
        
        # 生成PDF
        pdf_buffer = generate_declaration_pdf(
            declaration_text, 
            committee_name, 
            agenda, 
            participating_countries
        )
        
        # 将PDF转换为base64编码
        pdf_base64 = base64.b64encode(pdf_buffer.getvalue()).decode('utf-8')
        
        return jsonify({
            "success": True,
            "pdf_base64": pdf_base64,
            "filename": f"共同宣言_{agenda}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        })
        
    except Exception as e:
        print(f"导出PDF时出错: {str(e)}")
        return jsonify({"error": f"导出失败: {str(e)}"}), 500

def generate_declaration_pdf(declaration_text, committee_name, agenda, participating_countries):
    """生成共同宣言的PDF文件"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1,  # 居中
        textColor=colors.darkblue
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=20,
        alignment=1,  # 居中
        textColor=colors.darkblue
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12,
        leading=18
    )
    
    # 构建PDF内容
    story = []
    
    # 标题
    story.append(Paragraph(f"关于{agenda}的共同宣言", title_style))
    story.append(Spacer(1, 20))
    
    # 会议信息
    story.append(Paragraph(f"会议名称：{committee_name}", subtitle_style))
    story.append(Paragraph(f"议题：{agenda}", subtitle_style))
    story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}", subtitle_style))
    story.append(Spacer(1, 20))
    
    # 参与国家
    if participating_countries:
        story.append(Paragraph("参与国家：", normal_style))
        countries_text = "、".join(participating_countries)
        story.append(Paragraph(countries_text, normal_style))
        story.append(Spacer(1, 20))
    
    # 宣言内容
    story.append(Paragraph("共同宣言内容：", normal_style))
    
    # 将宣言文本分段处理
    paragraphs = declaration_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), normal_style))
            story.append(Spacer(1, 12))
    
    # 生成PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

@app.route('/api/get_declaration_history', methods=['GET'])
def get_declaration_history():
    """获取宣言历史记录"""
    try:
        session_id = request.args.get("session_id", "default")
        
        # 获取该会期的所有宣言记录
        declarations = list(db.declarations.find(
            {"session_id": session_id},
            sort=[("generated_at", -1)]
        ))
        
        # 格式化数据
        history = []
        for decl in declarations:
            history.append({
                "id": str(decl["_id"]),
                "declaration": decl.get("declaration", ""),
                "status": decl.get("status", "draft"),
                "participating_countries": decl.get("participating_countries", []),
                "generated_at": decl.get("generated_at"),
                "finalized_at": decl.get("finalized_at")
            })
        
        return jsonify({
            "success": True,
            "history": history
        })
        
    except Exception as e:
        print(f"获取宣言历史时出错: {str(e)}")
        return jsonify({"error": f"获取失败: {str(e)}"}), 500

# =========================
# 会议状态管理API
# =========================

@app.route('/api/meeting/state', methods=['GET'])
def api_get_meeting_state():
    """获取会议当前状态和可执行操作"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取会议设置
        meeting = cols["settings"].find_one({"session_id": session_id})
        if not meeting:
            return jsonify({
                "code": 404,
                "message": "会议不存在"
            }), 404
        
        # 获取当前阶段状态
        current_phase = meeting.get("meeting_state", {}).get("current_phase", "init")
        phase_locks = meeting.get("meeting_state", {}).get("phase_locks", {})
        chairman_controls = meeting.get("chairman_controls", {})
        
        # 根据当前阶段确定可用操作
        available_actions = []
        if current_phase == "init":
            available_actions = ["start_rollcall"]
        elif current_phase == "rollcall":
            available_actions = ["complete_rollcall", "pause_meeting"]
        elif current_phase == "file_submission":
            available_actions = ["start_motion", "pause_meeting"]
        elif current_phase == "motion":
            available_actions = ["start_voting", "pause_meeting"]
        elif current_phase == "voting":
            available_actions = ["generate_declaration", "pause_meeting"]
        elif current_phase == "declaration":
            available_actions = ["complete_meeting", "pause_meeting"]
        
        return jsonify({
            "code": 200,
            "data": {
                "current_phase": current_phase,
                "available_actions": available_actions,
                "phase_locks": phase_locks,
                "chairman_controls": chairman_controls,
                "phase_progress": {
                    "rollcall": 80 if current_phase in ["file_submission", "motion", "voting", "declaration", "completed"] else 0,
                    "file_submission": 60 if current_phase in ["motion", "voting", "declaration", "completed"] else 0,
                    "motion": 70 if current_phase in ["voting", "declaration", "completed"] else 0,
                    "voting": 90 if current_phase in ["declaration", "completed"] else 0,
                    "declaration": 100 if current_phase == "completed" else 0
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"获取会议状态失败: {str(e)}"
        }), 500

@app.route('/api/meeting/advance_phase', methods=['POST'])
def api_advance_meeting_phase():
    """主席手动推进到下一阶段"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        target_phase = data.get("target_phase", "")
        
        if not target_phase:
            return jsonify({
                "code": 400,
                "message": "目标阶段不能为空"
            }), 400
        
        cols = get_cols_by_session(session_id)
        
        # 获取当前会议状态
        meeting = cols["settings"].find_one({"session_id": session_id})
        if not meeting:
            return jsonify({
                "code": 404,
                "message": "会议不存在"
            }), 404
        
        current_phase = meeting.get("meeting_state", {}).get("current_phase", "init")
        
        # 检查阶段转换是否合法
        valid_transitions = {
            "init": ["rollcall"],
            "rollcall": ["file_submission"],
            "file_submission": ["motion"],
            "motion": ["voting"],
            "voting": ["declaration"],
            "declaration": ["completed"]
        }
        
        if target_phase not in valid_transitions.get(current_phase, []):
            return jsonify({
                "code": 400,
                "message": f"无法从 {current_phase} 阶段直接转换到 {target_phase} 阶段"
            }), 400
        
        # 更新会议状态
        current_time = datetime.now(UTC).isoformat() + "Z"
        
        # 完成当前阶段
        phase_history = meeting.get("meeting_state", {}).get("phase_history", [])
        for phase_record in phase_history:
            if phase_record["phase"] == current_phase and not phase_record.get("completed_at"):
                phase_record["completed_at"] = current_time
                break
        
        # 添加新阶段
        phase_history.append({
            "phase": target_phase,
            "started_at": current_time,
            "completed_at": None
        })
        
        # 更新数据库
        cols["settings"].update_one(
            {"session_id": session_id},
            {
                "$set": {
                    "meeting_state.current_phase": target_phase,
                    "meeting_state.phase_history": phase_history
                }
            }
        )
        
        return jsonify({
            "code": 200,
            "message": f"成功推进到 {target_phase} 阶段",
            "data": {
                "previous_phase": current_phase,
                "current_phase": target_phase,
                "transition_time": current_time
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"推进阶段失败: {str(e)}"
        }), 500

@app.route('/api/meeting/lock_phase', methods=['POST'])
def api_lock_meeting_phase():
    """锁定/解锁会议阶段"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default")
        phase = data.get("phase", "")
        locked = data.get("locked", False)
        
        if not phase:
            return jsonify({
                "code": 400,
                "message": "阶段名称不能为空"
            }), 400
        
        cols = get_cols_by_session(session_id)
        
        # 更新阶段锁定状态
        cols["settings"].update_one(
            {"session_id": session_id},
            {
                "$set": {
                    f"meeting_state.phase_locks.{phase}": locked
                }
            }
        )
        
        return jsonify({
            "code": 200,
            "message": f"阶段 {phase} 已{'锁定' if locked else '解锁'}",
            "data": {
                "phase": phase,
                "locked": locked
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"锁定阶段失败: {str(e)}"
        }), 500

@app.route('/api/meeting/phase_status', methods=['GET'])
def api_get_phase_status():
    """获取各阶段完成状态"""
    try:
        session_id = request.args.get("session_id", "default")
        cols = get_cols_by_session(session_id)
        
        # 获取点名状态
        rollcall_stats = cols["rollcall"].find_one({"session_id": session_id})
        arrived_countries = rollcall_stats.get("arrived_countries", []) if rollcall_stats else []
        
        # 获取文件提交状态
        submissions = list(cols["submissions"].find({"session_id": session_id}))
        submitted_countries = [sub["country_id"] for sub in submissions]
        
        # 获取投票状态
        vote_details = list(cols["file_vote_details"].find({"session_id": session_id}))
        voted_countries = list(set([vote["country_id"] for vote in vote_details]))
        
        return jsonify({
            "code": 200,
            "data": {
                "rollcall": {
                    "completed": len(arrived_countries) > 0,
                    "participants": len(arrived_countries),
                    "arrived_countries": arrived_countries
                },
                "file_submission": {
                    "completed": len(submitted_countries) > 0,
                    "submitted": len(submitted_countries),
                    "total": len(arrived_countries),
                    "submitted_countries": submitted_countries
                },
                "motion": {
                    "completed": False,  # 需要根据实际动议状态判断
                    "current_speaker": None
                },
                "voting": {
                    "completed": len(voted_countries) == len(arrived_countries) if arrived_countries else False,
                    "voted": len(voted_countries),
                    "total": len(arrived_countries)
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"获取阶段状态失败: {str(e)}"
        }), 500

# =========================
# Socket.io 实时通信
# =========================

@socketio.on('connect')
def handle_connect():
    """用户连接"""
    print(f"用户连接: {request.sid}")
    emit('connected', {'message': '连接成功'})

@socketio.on('disconnect')
def handle_disconnect():
    """用户断开连接"""
    print(f"用户断开连接: {request.sid}")
    # 这里可以更新用户状态为离线

@socketio.on('join_room')
def handle_join_room(data):
    """用户加入房间"""
    try:
        room_id = data.get('room_id')
        user_id = data.get('user_id')
        role = data.get('role', 'participant')
        
        if not room_id or not user_id:
            emit('error', {'message': '房间ID和用户ID不能为空'})
            return
        
        # 加入Socket.io房间
        join_room(room_id)
        
        # 广播用户加入消息
        emit('user_joined', {
            'user_id': user_id,
            'role': role,
            'message': f'用户已加入房间'
        }, room=room_id, include_self=False)
        
        # 发送房间当前状态
        emit('room_status', {
            'room_id': room_id,
            'message': '已加入房间'
        })
        
    except Exception as e:
        emit('error', {'message': f'加入房间失败: {str(e)}'})

@socketio.on('leave_room')
def handle_leave_room(data):
    """用户离开房间"""
    try:
        room_id = data.get('room_id')
        user_id = data.get('user_id')
        
        if not room_id or not user_id:
            emit('error', {'message': '房间ID和用户ID不能为空'})
            return
        
        # 离开Socket.io房间
        leave_room(room_id)
        
        # 广播用户离开消息
        emit('user_left', {
            'user_id': user_id,
            'message': f'用户已离开房间'
        }, room=room_id, include_self=False)
        
    except Exception as e:
        emit('error', {'message': f'离开房间失败: {str(e)}'})

@socketio.on('meeting_state_change')
def handle_meeting_state_change(data):
    """会议状态变更广播"""
    try:
        room_id = data.get('room_id')
        session_id = data.get('session_id')
        new_phase = data.get('new_phase')
        message = data.get('message', '会议状态已更新')
        
        if not room_id or not session_id:
            emit('error', {'message': '房间ID和会议ID不能为空'})
            return
        
        # 广播会议状态变更
        emit('meeting_phase_changed', {
            'session_id': session_id,
            'new_phase': new_phase,
            'message': message,
            'timestamp': datetime.now(UTC).isoformat() + "Z"
        }, room=room_id)
        
    except Exception as e:
        emit('error', {'message': f'广播会议状态失败: {str(e)}'})

@socketio.on('rollcall_update')
def handle_rollcall_update(data):
    """点名状态更新广播"""
    try:
        room_id = data.get('room_id')
        session_id = data.get('session_id')
        country_id = data.get('country_id')
        status = data.get('status')  # present, absent, pending
        
        if not room_id or not session_id:
            emit('error', {'message': '房间ID和会议ID不能为空'})
            return
        
        # 广播点名状态更新
        emit('rollcall_status_changed', {
            'session_id': session_id,
            'country_id': country_id,
            'status': status,
            'timestamp': datetime.now(UTC).isoformat() + "Z"
        }, room=room_id)
        
    except Exception as e:
        emit('error', {'message': f'广播点名状态失败: {str(e)}'})

@socketio.on('file_submission_update')
def handle_file_submission_update(data):
    """文件提交状态更新广播"""
    try:
        room_id = data.get('room_id')
        session_id = data.get('session_id')
        country_id = data.get('country_id')
        file_name = data.get('file_name')
        status = data.get('status')  # submitted, pending
        
        if not room_id or not session_id:
            emit('error', {'message': '房间ID和会议ID不能为空'})
            return
        
        # 广播文件提交状态更新
        emit('file_submission_changed', {
            'session_id': session_id,
            'country_id': country_id,
            'file_name': file_name,
            'status': status,
            'timestamp': datetime.now(UTC).isoformat() + "Z"
        }, room=room_id)
        
    except Exception as e:
        emit('error', {'message': f'广播文件提交状态失败: {str(e)}'})

@socketio.on('vote_update')
def handle_vote_update(data):
    """投票状态更新广播"""
    try:
        room_id = data.get('room_id')
        session_id = data.get('session_id')
        country_id = data.get('country_id')
        file_id = data.get('file_id')
        vote_result = data.get('vote_result')  # agree, disagree, abstain
        
        if not room_id or not session_id:
            emit('error', {'message': '房间ID和会议ID不能为空'})
            return
        
        # 广播投票状态更新
        emit('vote_status_changed', {
            'session_id': session_id,
            'country_id': country_id,
            'file_id': file_id,
            'vote_result': vote_result,
            'timestamp': datetime.now(UTC).isoformat() + "Z"
        }, room=room_id)
        
    except Exception as e:
        emit('error', {'message': f'广播投票状态失败: {str(e)}'})

# 注意：测试路由已移至独立的测试文件中

if __name__ == "__main__":
    app.config["TEMPLATES_AUTO_RELOAD"] = True
    app.jinja_env.auto_reload = True
    print("🚀 WTO模拟谈判系统启动中...")
    print("=" * 50)
    print("📍 系统接口地址：http://127.0.0.1:5000")
    print("👑 主席控制台：http://127.0.0.1:5000/chairman-selection")
    print("🌍 与会国门户：http://127.0.0.1:5000/country-portal")
    print("=" * 50)
    print("=" * 50)
    # 使用SocketIO启动应用 - 监听所有网络接口
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)
